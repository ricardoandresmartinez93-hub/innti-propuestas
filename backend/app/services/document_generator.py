"""
Servicio de generación de documentos Word y PDF.
Genera propuestas comerciales y anexos técnicos con la estructura estándar de Quipux.
"""
import re
from datetime import date
from html.parser import HTMLParser
from html import unescape
from pathlib import Path
from typing import List, Optional
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.services.portfolio_service import PortfolioProduct


# ---------------------------------------------------------------------------
# HTML → docx: helpers de bajo nivel
# ---------------------------------------------------------------------------

# Mapa color hex (aproximado) → highlight index de Word.
# WD_COLOR_INDEX solo soporta una paleta cerrada; mapeamos por familia.
_HEX_TO_HIGHLIGHT = {
    "ffff00": WD_COLOR_INDEX.YELLOW,
    "00ff00": WD_COLOR_INDEX.BRIGHT_GREEN,
    "00ffff": WD_COLOR_INDEX.TURQUOISE,
    "ff00ff": WD_COLOR_INDEX.PINK,
    "0000ff": WD_COLOR_INDEX.BLUE,
    "ff0000": WD_COLOR_INDEX.RED,
    "000080": WD_COLOR_INDEX.DARK_BLUE,
    "008080": WD_COLOR_INDEX.TEAL,
    "008000": WD_COLOR_INDEX.GREEN,
    "800080": WD_COLOR_INDEX.VIOLET,
    "800000": WD_COLOR_INDEX.DARK_RED,
    "808000": WD_COLOR_INDEX.DARK_YELLOW,
    "808080": WD_COLOR_INDEX.GRAY_50,
    "c0c0c0": WD_COLOR_INDEX.GRAY_25,
    "000000": WD_COLOR_INDEX.BLACK,
}

_ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _parse_style(style: str) -> dict:
    """Parsea una declaración CSS inline en un dict {prop: value}."""
    result: dict = {}
    if not style:
        return result
    for decl in style.split(";"):
        if ":" not in decl:
            continue
        prop, _, value = decl.partition(":")
        result[prop.strip().lower()] = value.strip()
    return result


def _normalize_hex(value: str) -> Optional[str]:
    """Normaliza colores `#rgb` / `#rrggbb` a 6 dígitos en minúsculas. Devuelve None si inválido."""
    if not value:
        return None
    value = value.strip().lstrip("#").lower()
    if re.fullmatch(r"[0-9a-f]{3}", value):
        value = "".join(c * 2 for c in value)
    if re.fullmatch(r"[0-9a-f]{6}", value):
        return value
    return None


def _closest_highlight(hex_color: str) -> WD_COLOR_INDEX:
    """Devuelve el highlight Word más cercano al hex dado (por distancia euclídea RGB)."""
    if hex_color in _HEX_TO_HIGHLIGHT:
        return _HEX_TO_HIGHLIGHT[hex_color]
    try:
        r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
    except ValueError:
        return WD_COLOR_INDEX.YELLOW
    best = WD_COLOR_INDEX.YELLOW
    best_dist = float("inf")
    for hx, idx in _HEX_TO_HIGHLIGHT.items():
        rr, gg, bb = int(hx[0:2], 16), int(hx[2:4], 16), int(hx[4:6], 16)
        dist = (r - rr) ** 2 + (g - gg) ** 2 + (b - bb) ** 2
        if dist < best_dist:
            best_dist = dist
            best = idx
    return best


def _add_hyperlink(paragraph, url: str, text: str, marks: dict) -> None:
    """Inserta un hyperlink en el párrafo aplicando las marcas activas.

    python-docx no expone hyperlinks de forma directa, por lo que se construye
    el elemento OOXML `w:hyperlink` manualmente.
    """
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    # Estilo "Hyperlink" si existe (azul + subrayado por defecto).
    style = OxmlElement("w:rStyle")
    style.set(qn("w:val"), "Hyperlink")
    rpr.append(style)

    # Color azul siempre (por encima de cualquier color textual del run).
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    rpr.append(color)

    # Marca de subrayado si la activa el usuario o por estilo del link.
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)

    if marks.get("bold"):
        rpr.append(OxmlElement("w:b"))
    if marks.get("italic"):
        rpr.append(OxmlElement("w:i"))
    if marks.get("strike"):
        rpr.append(OxmlElement("w:strike"))

    run.append(rpr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_horizontal_rule(doc: Document) -> None:
    """Inserta una línea horizontal (border-bottom en un párrafo vacío)."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ---------------------------------------------------------------------------
# Parser HTML → docx
# ---------------------------------------------------------------------------

_BLOCK_TAGS = {"p", "h1", "h2", "h3", "h4", "h5", "h6", "li", "blockquote", "pre"}
_INLINE_MARKS = {"strong", "b", "em", "i", "u", "s", "del", "sup", "sub", "mark", "code"}


class _HtmlToDocxParser(HTMLParser):
    """Recorre HTML de TipTap y emite contenido a un Document de python-docx.

    Mantiene una pila de marcas inline (`_marks_stack`) cuyas marcas se aplican
    al run actual. Los bloques (p, h*, li, blockquote, pre) abren un nuevo párrafo.
    """

    def __init__(self, doc: Document) -> None:
        super().__init__(convert_charrefs=False)
        self.doc = doc
        # Pila de marcas inline acumuladas. Cada entrada es un dict.
        self._marks_stack: List[dict] = []
        # Pila de contenedores de lista: cada entrada es "ul" o "ol".
        self._list_stack: List[str] = []
        # Párrafo en construcción (None entre bloques).
        self._current_paragraph = None
        # Estilo de párrafo pendiente para el siguiente bloque (p.ej. "Quote").
        self._pending_block_style: Optional[str] = None
        # Alineación pendiente para el siguiente párrafo.
        self._pending_alignment = None
        # Hyperlink activo (URL) si estamos dentro de un <a>.
        self._link_href: Optional[str] = None
        # Buffer de texto plano del run actual.
        self._text_buffer: str = ""

    # ----- helpers de marcas -------------------------------------------------

    def _current_marks(self) -> dict:
        """Combina la pila de marcas en un único dict."""
        merged: dict = {}
        for layer in self._marks_stack:
            merged.update({k: v for k, v in layer.items() if v is not None})
        return merged

    def _push_marks(self, marks: dict) -> None:
        self._flush_text()
        self._marks_stack.append(marks)

    def _pop_marks(self) -> None:
        self._flush_text()
        if self._marks_stack:
            self._marks_stack.pop()

    # ----- helpers de párrafo ------------------------------------------------

    def _ensure_paragraph(self) -> None:
        if self._current_paragraph is not None:
            return
        if self._list_stack and self._list_stack[-1] == "ol":
            style = "List Number"
        elif self._list_stack and self._list_stack[-1] == "ul":
            style = "List Bullet"
        else:
            style = self._pending_block_style or None
        self._current_paragraph = (
            self.doc.add_paragraph(style=style) if style else self.doc.add_paragraph()
        )
        if self._pending_alignment is not None:
            self._current_paragraph.alignment = self._pending_alignment
        self._pending_block_style = None
        self._pending_alignment = None

    def _close_paragraph(self) -> None:
        self._flush_text()
        self._current_paragraph = None

    def _flush_text(self) -> None:
        """Empuja el buffer de texto como un run con las marcas activas."""
        if not self._text_buffer:
            return
        text = self._text_buffer
        self._text_buffer = ""
        self._ensure_paragraph()
        marks = self._current_marks()
        if self._link_href:
            _add_hyperlink(self._current_paragraph, self._link_href, text, marks)
            return
        run = self._current_paragraph.add_run(text)
        if marks.get("bold"):
            run.bold = True
        if marks.get("italic"):
            run.italic = True
        if marks.get("underline"):
            run.underline = True
        if marks.get("strike"):
            run.font.strike = True
        if marks.get("superscript"):
            run.font.superscript = True
        if marks.get("subscript"):
            run.font.subscript = True
        if marks.get("code"):
            run.font.name = "Consolas"
        color_hex = marks.get("color")
        if color_hex:
            try:
                run.font.color.rgb = RGBColor.from_string(color_hex.upper())
            except (ValueError, AttributeError):
                pass
        highlight_hex = marks.get("highlight")
        if highlight_hex:
            run.font.highlight_color = _closest_highlight(highlight_hex)

    # ----- HTMLParser callbacks ---------------------------------------------

    def handle_starttag(self, tag: str, attrs):  # type: ignore[override]
        tag = tag.lower()
        attr_dict = {k.lower(): (v or "") for k, v in attrs}
        style = _parse_style(attr_dict.get("style", ""))

        if tag == "br":
            # Salto de línea dentro del párrafo actual.
            self._flush_text()
            self._ensure_paragraph()
            self._current_paragraph.add_run().add_break()
            return

        if tag == "hr":
            self._close_paragraph()
            _add_horizontal_rule(self.doc)
            return

        if tag in {"ul", "ol"}:
            self._close_paragraph()
            self._list_stack.append(tag)
            return

        if tag == "li":
            self._close_paragraph()
            # _ensure_paragraph aplicará el estilo correcto según _list_stack.
            self._ensure_paragraph()
            return

        if tag in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            self._close_paragraph()
            level = int(tag[1])
            self._current_paragraph = self.doc.add_heading("", level=level)
            align = style.get("text-align")
            if align in _ALIGN_MAP:
                self._current_paragraph.alignment = _ALIGN_MAP[align]
            return

        if tag == "p":
            self._close_paragraph()
            align = style.get("text-align")
            if align in _ALIGN_MAP:
                self._pending_alignment = _ALIGN_MAP[align]
            return

        if tag == "blockquote":
            self._close_paragraph()
            self._pending_block_style = "Quote"
            return

        if tag == "pre":
            self._close_paragraph()
            self._pending_block_style = "No Spacing"
            self._push_marks({"code": True})
            return

        if tag == "a":
            self._flush_text()
            self._link_href = attr_dict.get("href") or None
            self._push_marks({})  # capa vacía para que pop simétrico funcione
            return

        # ----- marcas inline -----
        if tag in {"strong", "b"}:
            self._push_marks({"bold": True})
            return
        if tag in {"em", "i"}:
            self._push_marks({"italic": True})
            return
        if tag == "u":
            self._push_marks({"underline": True})
            return
        if tag in {"s", "del"}:
            self._push_marks({"strike": True})
            return
        if tag == "sup":
            self._push_marks({"superscript": True})
            return
        if tag == "sub":
            self._push_marks({"subscript": True})
            return
        if tag == "code":
            self._push_marks({"code": True})
            return
        if tag == "mark":
            hex_color = _normalize_hex(style.get("background-color", "ffff00"))
            self._push_marks({"highlight": hex_color or "ffff00"})
            return
        if tag == "span":
            layer: dict = {}
            color = _normalize_hex(style.get("color", ""))
            if color:
                layer["color"] = color
            bg = _normalize_hex(style.get("background-color", ""))
            if bg:
                layer["highlight"] = bg
            self._push_marks(layer)
            return

        # Tags desconocidos: empujar capa vacía para mantener simetría con endtag.
        self._push_marks({})

    def handle_endtag(self, tag: str):  # type: ignore[override]
        tag = tag.lower()
        if tag in {"br", "hr"}:
            return

        if tag in {"ul", "ol"}:
            self._close_paragraph()
            if self._list_stack:
                self._list_stack.pop()
            return

        if tag == "li":
            self._close_paragraph()
            return

        if tag in {"p", "blockquote"} or tag in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            self._close_paragraph()
            return

        if tag == "pre":
            self._pop_marks()  # cierra la marca "code" abierta
            self._close_paragraph()
            return

        if tag == "a":
            self._flush_text()
            self._link_href = None
            self._pop_marks()
            return

        # Marcas inline conocidas y desconocidas: pop simétrico
        if tag in _INLINE_MARKS or tag == "span":
            self._pop_marks()
            return

        self._pop_marks()

    def handle_startendtag(self, tag: str, attrs):  # type: ignore[override]
        # <br/>, <hr/> auto-cerrados
        tag = tag.lower()
        if tag == "br":
            self._flush_text()
            self._ensure_paragraph()
            self._current_paragraph.add_run().add_break()
            return
        if tag == "hr":
            self._close_paragraph()
            _add_horizontal_rule(self.doc)
            return
        # Otros auto-cerrados: tratar como start+end vacíos
        self.handle_starttag(tag, attrs)
        self.handle_endtag(tag)

    def handle_data(self, data: str):  # type: ignore[override]
        self._text_buffer += data

    def handle_entityref(self, name: str):  # type: ignore[override]
        self._text_buffer += unescape(f"&{name};")

    def handle_charref(self, name: str):  # type: ignore[override]
        self._text_buffer += unescape(f"&#{name};")

    def finalize(self) -> None:
        self._close_paragraph()


def _add_html_paragraphs(doc: Document, html_text: str) -> None:
    """Convierte HTML de TipTap en contenido Word preservando formato.

    Las marcas de TipTap (negrita, cursiva, subrayado, tachado, super/subíndice,
    color, resaltado, enlaces, alineación, citas, líneas horizontales y listas)
    se mapean a runs/párrafos de python-docx. Texto sin etiquetas se inserta
    como un único párrafo.
    """
    if not html_text:
        return

    # Si el contenido no contiene tags, insertarlo como párrafo plano
    # (caso común de campos legacy en la BD).
    if "<" not in html_text:
        plain = html_text.strip()
        if plain:
            doc.add_paragraph(plain)
        return

    parser = _HtmlToDocxParser(doc)
    parser.feed(html_text)
    parser.finalize()


def _has_content(html: Optional[str]) -> bool:
    """
    Devuelve True si el HTML tiene contenido real (no solo tags vacíos/whitespace).
    Se usa para decidir si incluir una sección en el documento generado.
    """
    if not html:
        return False
    stripped = _strip_html(html).strip()
    return len(stripped) > 0


def _strip_html(text: str) -> str:
    """Elimina etiquetas HTML y decodifica entidades básicas para inserción en Word.

    TipTap guarda HTML en la BD (<p>, <strong>, <ul>, etc.).
    Esta función produce texto plano apto para doc.add_paragraph().
    """
    # Convertir <br> y bloques de cierre en saltos de línea para no perder separación
    text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'</p>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'</li>', '\n', text, flags=re.IGNORECASE)
    # Eliminar todas las demás etiquetas
    text = re.sub(r'<[^>]+>', '', text)
    # Decodificar entidades HTML básicas
    text = (
        text.replace('&nbsp;', ' ')
            .replace('&amp;', '&')
            .replace('&lt;', '<')
            .replace('&gt;', '>')
            .replace('&quot;', '"')
            .replace('&#39;', "'")
    )
    # Normalizar espacios en blanco sin destruir saltos de línea intencionales
    lines = [re.sub(r'[ \t]+', ' ', line).strip() for line in text.split('\n')]
    return '\n'.join(line for line in lines if line)


# Mapa de meses en español para fecha de carta
_MONTHS_ES = {
    1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
    5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
    9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre",
}


class DocumentGeneratorError(Exception):
    """Error en la generación de documentos."""
    pass


class DocumentGenerator:
    """Genera documentos Word con la estructura estándar de propuestas Quipux."""

    # ------------------------------------------------------------------ #
    # Textos fijos (basados en propuestas reales de la empresa)            #
    # ------------------------------------------------------------------ #

    # Servicios excluidos comunes a contratos con entrega de software (licensing y support_maintenance).
    # SaaS (services) no excluye nada por defecto: el modelo incluye infraestructura y operación.
    _EXCLUDED_LICENSING_AND_SUPPORT = [
        "Infraestructura tecnológica centralizada (su suministro y mantenimiento es responsabilidad del cliente).",
        "Desarrollo e implantación de nuevas funcionalidades no incluidas en el alcance de la propuesta.",
        "Solución a dudas sobre la operación de sistemas diferentes a los ofertados en la presente propuesta.",
        "Actualización de versión del motor de base de datos y su respectivo mantenimiento.",
        "Procesos de configuración de redes, servidores, equipos, impresoras, sistemas operativos y comunicaciones.",
        "Interrogantes sobre el modelo de datos, diccionario de datos y características de programación de las "
        "soluciones (considerados confidenciales).",
        "Suministro de código fuente, modelos de datos o documentación técnica considerada confidencial.",
        "Asesoría para la solución de errores o mala operación provocados por factores externos al alcance del contrato.",
        "Corrección de fallas originadas por defectos de hardware, redes o instalaciones locativas "
        "(responsabilidad del cliente).",
        "Realizaciones de migraciones, cruces y depuraciones de información de bases de datos u otras "
        "actividades análogas.",
    ]

    # Servicios excluidos por esquema. Para SaaS (services) la lista es vacía por defecto
    # — la propuesta incluye infraestructura, operación y soporte, así que no hay nada que excluir.
    # Decisión del PDF de reunión: "esto no va, por ejemplo, en las que son de software as a service".
    EXCLUDED_SERVICES_BY_SCHEME: dict = {
        "licensing": _EXCLUDED_LICENSING_AND_SUPPORT,
        "services": [],
        "support_maintenance": _EXCLUDED_LICENSING_AND_SUPPORT,
    }

    # Lista por defecto (retrocompatibilidad y fallback): equivale a licensing.
    EXCLUDED_SERVICES = _EXCLUDED_LICENSING_AND_SUPPORT

    # Propiedad intelectual por esquema. La transcripción de la reunión es explícita:
    # "este cambia, el primero, la propiedad intelectual cambia dependiendo del esquema".
    # Los textos exactos para services y support_maintenance deben validarse con jurídica
    # — aquí dejamos defaults razonables basados en el texto de licenciamiento.
    # {client_entity} se reemplaza con el nombre de la entidad cliente.
    IP_TEXT_BY_SCHEME: dict = {
        "licensing": (
            "La arquitectura, los diseños técnicos, comerciales, económicos, financieros y administrativos que "
            "QUIPUX realice, y en general el Know How asociado, son de su plena y exclusiva propiedad. QUIPUX puede "
            "utilizarlos libremente, introducir cambios, modificaciones y explotarlos económicamente en cualquier parte "
            "del mundo. {client_entity} no será propietario de ninguna clase de derechos de autor ni de Propiedad "
            "Industrial sobre las soluciones de QUIPUX, quien podrá dar la utilización que encuentre conveniente, "
            "diferente a la que pueda ser autorizada expresamente en el presente documento y los Anexos."
        ),
        "services": (
            "Los servicios prestados por QUIPUX a {client_entity} bajo este esquema utilizan plataformas, "
            "metodologías y componentes propiedad de QUIPUX. El uso de los servicios no transfiere a {client_entity} "
            "derechos de autor, marca, ni propiedad industrial sobre dichas plataformas o componentes. "
            "Los entregables del servicio (reportes, configuraciones, parametrizaciones) son de uso del cliente "
            "para los fines del contrato, sin perjuicio de los derechos de QUIPUX sobre las herramientas subyacentes."
        ),
        "support_maintenance": (
            "Las actualizaciones, parches, mejoras y nuevas versiones que QUIPUX desarrolle como parte del soporte "
            "y mantenimiento son de su exclusiva propiedad intelectual. {client_entity} obtiene el derecho de uso "
            "de dichas mejoras durante la vigencia del contrato, sin que ello implique transferencia de propiedad "
            "industrial ni de derechos de autor. La documentación técnica entregada conserva su carácter confidencial."
        ),
    }

    # Texto por defecto (fallback cuando el esquema no está mapeado): equivale a licensing.
    IP_TEXT = IP_TEXT_BY_SCHEME["licensing"]

    @classmethod
    def get_ip_text(cls, scheme_type: str) -> str:
        """Devuelve el texto de propiedad intelectual aplicable al esquema dado."""
        return cls.IP_TEXT_BY_SCHEME.get(scheme_type.lower(), cls.IP_TEXT)

    @classmethod
    def get_excluded_services(cls, scheme_type: str) -> list:
        """Devuelve la lista de servicios excluidos aplicable al esquema dado.

        SaaS (services) devuelve lista vacía, indicando que no debe renderizarse
        la sección de servicios excluidos para ese esquema.
        """
        return cls.EXCLUDED_SERVICES_BY_SCHEME.get(scheme_type.lower(), cls.EXCLUDED_SERVICES)

    # Confidencialidad — {client_entity} se reemplaza con el nombre de la entidad cliente
    CONFIDENTIALITY_TEXT = (
        "La información no deberá ser divulgada fuera de {client_entity}, ni deberá ser duplicada, usada o "
        "divulgada para propósito diferente al de la evaluación de la presente propuesta. La información contenida "
        "en este documento constituye secretos de marca del cliente e información financiera y comercial de QUIPUX "
        "considerada como confidencial."
    )

    # Principios de prevención de actividades delictivas
    # {client_entity} se reemplaza con el nombre de la entidad cliente
    CRIME_PREVENTION_TEXT = (
        "{client_entity} declara que a la fecha de la presente propuesta no existen antecedentes de sanciones o "
        "investigaciones por fraude, soborno, corrupción u otras actividades delictivas. Se compromete a que en el "
        "desarrollo y ejecución del contrato adoptará los procedimientos de detección y prevención de actividades "
        "irregulares y de corrupción necesarios, y reportará inmediatamente a QUIPUX cualquier sospecha de estos "
        "actos, prestando toda la colaboración necesaria para garantizar el normal desarrollo del contrato."
    )

    # Transparencia y ética (referencia breve a la línea ética)
    ETHICS_TEXT = (
        "Para el cumplimiento al Programa de Transparencia y Ética Empresarial de QUIPUX, la línea ética de "
        "QUIPUX es: lineaetica@quipux.com | www.quipux.com"
    )

    # Nota de indexación IPC — solo para esquemas services y support_maintenance
    IPC_INDEXATION_TEXT = (
        "Los valores se indexarán cada año según el IPC de cierre del año anterior certificado por el DANE. "
        "Para los servicios en los que exista prestación de personal, el valor también se indexará de acuerdo "
        "con el incremento del salario mínimo."
    )

    # Plazo por defecto cuando no se provee texto personalizado
    DEFAULT_VALIDITY_TEXT = (
        "La vigencia del contrato será desde la fecha de suscripción hasta la terminación del mismo, "
        "conforme a lo acordado entre las partes."
    )

    def _add_table_of_contents(self, doc: Document) -> None:
        """
        Inserta un campo TOC real usando OOXML field codes.

        El atributo ``w:dirty="true"`` le indica a Word que regenere el índice
        automáticamente al abrir el documento.  Sin necesidad de intervención
        manual: Word detecta que el campo está desactualizado y lo reconstruye.

        Instrucción usada: ``TOC \\o "1-3" \\h \\z \\u``
          - ``\\o "1-3"`` → captura Heading 1, 2 y 3
          - ``\\h``        → crea hipervínculos internos en cada entrada
          - ``\\z``        → oculta numeración en vista web
          - ``\\u``        → usa el nivel de esquema del párrafo
        """
        paragraph = doc.add_paragraph()
        paragraph.style = doc.styles["Normal"]
        paragraph.paragraph_format.space_after = Pt(0)

        # Run 1 — BEGIN del campo, marcado como sucio (dirty) para forzar actualización
        r = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        fldChar.set(qn('w:dirty'), 'true')
        r._r.append(fldChar)

        # Run 2 — instrucción del campo TOC
        r = paragraph.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
        r._r.append(instrText)

        # Run 3 — SEPARATE (divide la instrucción del contenido calculado)
        r = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'separate')
        r._r.append(fldChar)

        # Run 4 — END del campo
        r = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        r._r.append(fldChar)

    def _add_page_number(self, paragraph):
        """Añade numeración de página al párrafo."""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar)

        run = paragraph.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        run._r.append(instrText)

        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar)

    @classmethod
    def get_payment_type_from_scheme(cls, scheme_type: str) -> str:
        """Mapea el tipo de esquema a un tipo de pago en español."""
        mapping = {
            "licensing": "pago único",
            "services": "pago mensual",
            "support_maintenance": "pago anual"
        }
        return mapping.get(scheme_type.lower(), "por definir")

    @classmethod
    def get_value_column_label(cls, scheme_type: str) -> str:
        """Devuelve el encabezado de la columna de valor según el esquema."""
        mapping = {
            "licensing": "Valor (IVA incluido)",
            "services": "Valor mensual (IVA incluido)",
            "support_maintenance": "Valor anual (IVA incluido)",
        }
        return mapping.get(scheme_type.lower(), "Valor")

    def _add_economic_conditions_table(
        self,
        doc: Document,
        scheme_type: str,
        products: List[PortfolioProduct],
    ) -> None:
        """
        Agrega una tabla de condiciones económicas con el formato estándar de Quipux.
        Genera una fila por producto/componente real más las filas de totales.
        Agrega nota de indexación IPC para esquemas de servicios y mantenimiento.
        """
        value_label = self.get_value_column_label(scheme_type)

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        # Encabezado (Azul Quipux #1a365d, texto blanco)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Componente'
        hdr_cells[1].text = value_label

        for cell in hdr_cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), '1A365D')
            tcPr.append(shd)

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # Filas dinámicas: una por producto (o 3 placeholders si no hay productos)
        if products:
            for product in products:
                row_cells = table.add_row().cells
                row_cells[0].text = product.name
                row_cells[1].text = '[Por definir]'
        else:
            for _ in range(3):
                row_cells = table.add_row().cells
                row_cells[0].text = '[Componente]'
                row_cells[1].text = '[Por definir]'

        # Filas de totales
        totals = [
            ("Subtotal", "[Por definir]"),
            ("IVA (19%)", "[Por definir]"),
            ("Total", "[Por definir]"),
        ]
        for label, val in totals:
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = val
            row_cells[0].paragraphs[0].runs[0].font.bold = True

        # Nota de indexación IPC para servicios y mantenimiento
        if scheme_type.lower() in ("services", "support_maintenance"):
            p_ipc = doc.add_paragraph()
            run_ipc = p_ipc.add_run(self.IPC_INDEXATION_TEXT)
            run_ipc.italic = True
            run_ipc.font.size = Pt(9)

    def generate_proposal_docx(
        self,
        title: str,
        client_name: str,
        client_position: str,
        client_entity: str,
        client_city: str,
        scheme_types: List[str],
        products: List[PortfolioProduct],
        context_text: str,
        scope_text: str,
        letter_text: str,
        validity_period: Optional[str] = None,
        economic_conditions: Optional[str] = None,
        payment_terms: Optional[str] = None,
        payment_frequency: Optional[str] = None,
        excluded_services: Optional[str] = None,
        ip_section: Optional[str] = None,
    ) -> Document:
        """
        Genera el documento Word de la propuesta comercial con formato profesional Quipux.

        Returns:
            Objeto Document de python-docx listo para guardar.
        """
        doc = Document()

        # --- Configuración de Márgenes ---
        section = doc.sections[0]
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

        # --- Configurar estilos ---
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.space_after = Pt(10)

        h1 = doc.styles["Heading 1"]
        h1.font.name = "Calibri"
        h1.font.size = Pt(16)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        h1.paragraph_format.space_before = Pt(18)
        h1.paragraph_format.space_after = Pt(12)

        h2 = doc.styles["Heading 2"]
        h2.font.name = "Calibri"
        h2.font.size = Pt(13)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
        h2.paragraph_format.space_before = Pt(14)
        h2.paragraph_format.space_after = Pt(8)

        # --- PORTADA ---
        doc.add_paragraph("\n\n\n")

        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_logo = p_logo.add_run("[LOGO QUIPUX]")
        run_logo.bold = True
        run_logo.font.size = Pt(12)

        doc.add_paragraph("\n\n")

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("PROPUESTA COMERCIAL")
        run.bold = True
        run.font.size = Pt(26)
        run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title.upper())
        run.font.size = Pt(18)
        run.bold = True

        doc.add_paragraph("\n\n\n")

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Presentada a:\n{client_entity}")
        run.font.size = Pt(14)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"\n{client_name}\n{client_position}")
        run.font.size = Pt(12)

        doc.add_page_break()

        # --- TABLA DE CONTENIDO ---
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("TABLA DE CONTENIDO")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        doc.add_paragraph("")
        self._add_table_of_contents(doc)

        doc.add_page_break()

        # --- CARTA DE PRESENTACIÓN ---
        today = date.today()
        fecha_str = f"{today.day} de {_MONTHS_ES[today.month]} de {today.year}"

        doc.add_paragraph(f"{client_city}, {fecha_str}")
        doc.add_paragraph("")
        p = doc.add_paragraph("Señor(a)")
        p.paragraph_format.space_after = Pt(0)
        p = doc.add_paragraph(f"{client_name}")
        p.paragraph_format.space_after = Pt(0)
        p = doc.add_paragraph(f"{client_position}")
        p.paragraph_format.space_after = Pt(0)
        p = doc.add_paragraph(f"{client_entity}")

        doc.add_paragraph("")
        p = doc.add_paragraph()
        run = p.add_run(f"Asunto: {title}")
        run.bold = True
        doc.add_paragraph("")

        if letter_text:
            _add_html_paragraphs(doc, letter_text)
        else:
            doc.add_paragraph(
                "Estamos presentando la propuesta comercial descrita en el asunto. "
                "Esperamos que esté ajustada a las expectativas y necesidades del proyecto."
            )

        doc.add_paragraph("")
        doc.add_paragraph("Atentamente,")
        doc.add_paragraph("")
        p = doc.add_paragraph()
        run = p.add_run("Juan Pablo Ramírez Madrid")
        run.bold = True
        p = doc.add_paragraph("Vicepresidente de Nuevos Negocios")
        p.paragraph_format.space_after = Pt(0)
        doc.add_paragraph("Quipux S.A.S.")

        doc.add_page_break()

        # ------------------------------------------------------------------ #
        # CUERPO DEL DOCUMENTO                                                 #
        # ------------------------------------------------------------------ #
        
        section_num = 1

        # 1. CONTEXTO
        if _has_content(context_text):
            doc.add_heading(f"{section_num}. CONTEXTO", level=1)
            _add_html_paragraphs(doc, context_text)
            section_num += 1

        # 2. ALCANCE GENERAL DE LA PROPUESTA
        if _has_content(scope_text) or products:
            doc.add_heading(f"{section_num}. ALCANCE GENERAL DE LA PROPUESTA", level=1)
            if _has_content(scope_text):
                _add_html_paragraphs(doc, scope_text)
            
            if products:
                doc.add_paragraph(
                    "El alcance incluye las siguientes soluciones y/o componentes:"
                )
                # Agrupar por categoría si los productos tienen ese campo
                categories: dict = {}
                for product in products:
                    cat = (product.category or "").strip()
                    categories.setdefault(cat, []).append(product)

                if len(categories) == 1 and "" in categories:
                    # Sin categorías — lista plana
                    for product in products:
                        doc.add_paragraph(product.name, style="List Bullet")
                else:
                    # Con categorías — subsecciones
                    for cat, cat_products in categories.items():
                        if cat:
                            h2_cat = doc.add_heading(cat, level=2)
                        for product in cat_products:
                            doc.add_paragraph(product.name, style="List Bullet")

            doc.add_paragraph(
                "Nota: Invitamos a leer los anexos técnicos para identificar y comprender "
                "el alcance de cada uno de los componentes relacionados en la propuesta comercial."
            )
            section_num += 1

        # 3. PLAZO
        if _has_content(validity_period):
            doc.add_heading(f"{section_num}. PLAZO", level=1)
            _add_html_paragraphs(doc, validity_period)
            section_num += 1

        # 4. CONDICIONES ECONÓMICAS
        if _has_content(economic_conditions) or _has_content(payment_terms) or products:
            doc.add_heading(f"{section_num}. CONDICIONES ECONÓMICAS", level=1)
            if _has_content(economic_conditions):
                _add_html_paragraphs(doc, economic_conditions)
            else:
                primary_scheme = scheme_types[0] if scheme_types else "licensing"
                self._add_economic_conditions_table(doc, primary_scheme, products)

            if _has_content(payment_terms):
                doc.add_heading("Facturación y forma de pago", level=2)
                _add_html_paragraphs(doc, payment_terms)
            section_num += 1

        # 5. SERVICIOS EXCLUIDOS
        if _has_content(excluded_services):
            doc.add_heading(f"{section_num}. SERVICIOS EXCLUIDOS", level=1)
            _add_html_paragraphs(doc, excluded_services)
            section_num += 1
        elif excluded_services is None:
            # Fallback para casos donde se quiera mantener el comportamiento original si no se especifica
            doc.add_heading(f"{section_num}. SERVICIOS EXCLUIDOS", level=1)
            doc.add_paragraph("La presente propuesta no incluye los siguientes servicios:")
            for item in self.EXCLUDED_SERVICES:
                doc.add_paragraph(item, style="List Bullet")
            section_num += 1
        else:
             # Si es string vacío explícito, no incluimos la sección según Tarea 2
             pass

        # 6. ESQUEMA DE PRESTACIÓN DE SERVICIOS Y LICENCIAMIENTO
        # Esta sección contiene subsecciones (6.1, 6.2, 6.3, 6.4).
        # La 6.1 (IP) es opcional según la tarea 2.
        # Las otras (6.2, 6.3, 6.4) parecen estructurales pero la 6.1 solo se incluye si tiene contenido.
        
        doc.add_heading(f"{section_num}. ESQUEMA DE PRESTACIÓN DE SERVICIOS Y LICENCIAMIENTO", level=1)

        # 6.1 Propiedad Intelectual
        if _has_content(ip_section):
            doc.add_heading(f"{section_num}.1. Propiedad Intelectual", level=2)
            _add_html_paragraphs(doc, ip_section)
        elif not ip_section and ip_section is not None:
            # String vacío explícito -> omitir
            pass
        else:
            # Fallback original
            doc.add_heading(f"{section_num}.1. Propiedad Intelectual", level=2)
            doc.add_paragraph(self.IP_TEXT.format(client_entity=client_entity))

        # 6.2 Confidencialidad
        doc.add_heading(f"{section_num}.2. Confidencialidad", level=2)
        doc.add_paragraph(self.CONFIDENTIALITY_TEXT.format(client_entity=client_entity))

        # 6.3 Principios de Prevención de Actividades Delictivas
        doc.add_heading(f"{section_num}.3. Principios de Prevención de Actividades Delictivas", level=2)
        doc.add_paragraph(self.CRIME_PREVENTION_TEXT.format(client_entity=client_entity))

        # 6.4 Cumplimiento al Programa de Transparencia y Ética Empresarial de Quipux
        doc.add_heading(
            f"{section_num}.4. Cumplimiento al Programa de Transparencia y Ética Empresarial de Quipux",
            level=2,
        )
        doc.add_paragraph(self.ETHICS_TEXT)


        # --- PIE DE PÁGINA (Numeración) ---
        footer = section.footer
        p_footer = footer.paragraphs[0]
        self._add_page_number(p_footer)

        return doc

    SCHEME_HEADING = {
        "licensing": "LICENCIAMIENTO",
        "services": "PRESTACIÓN DE SERVICIOS",
        "support_maintenance": "SOPORTE Y MANTENIMIENTO",
    }

    def _render_header_and_letter(
        self,
        doc: Document,
        title: str,
        client_name: str,
        client_position: str,
        client_entity: str,
        client_city: str,
        letter_text: str,
    ) -> None:
        """Portada, tabla de contenido y carta de presentación. Una sola vez por documento."""
        # --- PORTADA ---
        doc.add_paragraph("\n\n\n")
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_logo = p_logo.add_run("[LOGO QUIPUX]")
        run_logo.bold = True
        run_logo.font.size = Pt(12)

        doc.add_paragraph("\n\n")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("PROPUESTA COMERCIAL")
        run.bold = True
        run.font.size = Pt(26)
        run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title.upper())
        run.font.size = Pt(18)
        run.bold = True

        doc.add_paragraph("\n\n\n")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Presentada a:\n{client_entity}")
        run.font.size = Pt(14)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"\n{client_name}\n{client_position}")
        run.font.size = Pt(12)
        doc.add_page_break()

        # --- TABLA DE CONTENIDO ---
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("TABLA DE CONTENIDO")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        doc.add_paragraph("")
        self._add_table_of_contents(doc)
        doc.add_page_break()

        # --- CARTA DE PRESENTACIÓN ---
        today = date.today()
        fecha_str = f"{today.day} de {_MONTHS_ES[today.month]} de {today.year}"
        doc.add_paragraph(f"{client_city}, {fecha_str}")
        doc.add_paragraph("")
        p = doc.add_paragraph("Señor(a)")
        p.paragraph_format.space_after = Pt(0)
        p = doc.add_paragraph(f"{client_name}")
        p.paragraph_format.space_after = Pt(0)
        p = doc.add_paragraph(f"{client_position}")
        p.paragraph_format.space_after = Pt(0)
        p = doc.add_paragraph(f"{client_entity}")
        doc.add_paragraph("")
        p = doc.add_paragraph()
        run = p.add_run(f"Asunto: {title}")
        run.bold = True
        doc.add_paragraph("")

        if letter_text:
            _add_html_paragraphs(doc, letter_text)
        else:
            doc.add_paragraph(
                "Estamos presentando la propuesta comercial descrita en el asunto. "
                "Esperamos que esté ajustada a las expectativas y necesidades del proyecto."
            )

        doc.add_paragraph("")
        doc.add_paragraph("Atentamente,")
        doc.add_paragraph("")
        p = doc.add_paragraph()
        run = p.add_run("Juan Pablo Ramírez Madrid")
        run.bold = True
        p = doc.add_paragraph("Vicepresidente de Nuevos Negocios")
        p.paragraph_format.space_after = Pt(0)
        doc.add_paragraph("Quipux S.A.S.")
        doc.add_page_break()

    def _render_scheme_block(
        self,
        doc: Document,
        section_num: int,
        scheme_payload: dict,
        products: List[PortfolioProduct],
        scheme_heading: Optional[str] = None,
    ) -> None:
        """Renderiza el bloque de secciones de UN esquema (alcance, plazo, condiciones, exclusiones, IP).

        Si scheme_heading se provee, agrega un encabezado "N. ESQUEMA: {nombre}" y las
        subsecciones se numeran como N.1, N.2, etc. Si no, usa numeración plana N, N+1...
        """
        scheme_type = scheme_payload.get("scheme_type", "licensing")
        scope_text = scheme_payload.get("scope_text", "")
        validity_period = scheme_payload.get("validity_period")
        economic_conditions = scheme_payload.get("economic_conditions")
        payment_terms = scheme_payload.get("payment_terms")
        excluded_services = scheme_payload.get("excluded_services_text")
        ip_section = scheme_payload.get("ip_section_text")

        if scheme_heading:
            doc.add_heading(f"{section_num}. ESQUEMA: {scheme_heading}", level=1)
            sub = lambda i: f"{section_num}.{i}"
            heading_level = 2
        else:
            sub = lambda i: f"{section_num + i - 1}"
            heading_level = 1

        sub_idx = 1

        # Alcance
        if _has_content(scope_text) or products:
            doc.add_heading(f"{sub(sub_idx)}. ALCANCE GENERAL DE LA PROPUESTA", level=heading_level)
            if _has_content(scope_text):
                _add_html_paragraphs(doc, scope_text)
            if products:
                doc.add_paragraph("El alcance incluye las siguientes soluciones y/o componentes:")
                categories: dict = {}
                for product in products:
                    cat = (product.category or "").strip()
                    categories.setdefault(cat, []).append(product)
                if len(categories) == 1 and "" in categories:
                    for product in products:
                        doc.add_paragraph(product.name, style="List Bullet")
                else:
                    for cat, cat_products in categories.items():
                        if cat:
                            doc.add_heading(cat, level=heading_level + 1)
                        for product in cat_products:
                            doc.add_paragraph(product.name, style="List Bullet")
            doc.add_paragraph(
                "Nota: Invitamos a leer los anexos técnicos para identificar y comprender "
                "el alcance de cada uno de los componentes relacionados en la propuesta comercial."
            )
            sub_idx += 1

        # Plazo
        if _has_content(validity_period):
            doc.add_heading(f"{sub(sub_idx)}. PLAZO", level=heading_level)
            _add_html_paragraphs(doc, validity_period)
            sub_idx += 1

        # Condiciones económicas
        if _has_content(economic_conditions) or _has_content(payment_terms) or products:
            doc.add_heading(f"{sub(sub_idx)}. CONDICIONES ECONÓMICAS", level=heading_level)
            if _has_content(economic_conditions):
                _add_html_paragraphs(doc, economic_conditions)
            else:
                self._add_economic_conditions_table(doc, scheme_type, products)
            if _has_content(payment_terms):
                doc.add_heading("Facturación y forma de pago", level=heading_level + 1)
                _add_html_paragraphs(doc, payment_terms)
            sub_idx += 1

        # Servicios excluidos (omitir si es string vacío explícito — caso SaaS)
        if _has_content(excluded_services):
            doc.add_heading(f"{sub(sub_idx)}. SERVICIOS EXCLUIDOS", level=heading_level)
            _add_html_paragraphs(doc, excluded_services)
            sub_idx += 1

        # Propiedad intelectual del esquema (omitir si es string vacío explícito)
        if _has_content(ip_section):
            doc.add_heading(f"{sub(sub_idx)}. PROPIEDAD INTELECTUAL", level=heading_level)
            _add_html_paragraphs(doc, ip_section)
            sub_idx += 1

    def _render_legal_footer(self, doc: Document, section_num: int, client_entity: str) -> None:
        """Aspectos legales que NO dependen del esquema (confidencialidad, prevención, ética)."""
        doc.add_heading(f"{section_num}. ASPECTOS LEGALES", level=1)
        doc.add_heading(f"{section_num}.1. Confidencialidad", level=2)
        doc.add_paragraph(self.CONFIDENTIALITY_TEXT.format(client_entity=client_entity))
        doc.add_heading(f"{section_num}.2. Principios de Prevención de Actividades Delictivas", level=2)
        doc.add_paragraph(self.CRIME_PREVENTION_TEXT.format(client_entity=client_entity))
        doc.add_heading(
            f"{section_num}.3. Cumplimiento al Programa de Transparencia y Ética Empresarial de Quipux",
            level=2,
        )
        doc.add_paragraph(self.ETHICS_TEXT)

    def generate_combined_proposal_docx(
        self,
        title: str,
        client_name: str,
        client_position: str,
        client_entity: str,
        client_city: str,
        products: List[PortfolioProduct],
        context_text: str,
        letter_text: str,
        schemes_payload: List[dict],
    ) -> Document:
        """Genera un único Word combinando N esquemas en bloques separados.

        Estructura:
            Portada + Índice + Carta (una vez)
            1. CONTEXTO (global)
            2. ESQUEMA: LICENCIAMIENTO
               2.1 Alcance, 2.2 Plazo, 2.3 Condiciones económicas, 2.4 Exclusiones, 2.5 IP
            3. ESQUEMA: PRESTACIÓN DE SERVICIOS
               3.1 ...
            ...
            N. ASPECTOS LEGALES (confidencialidad, prevención, ética)
        """
        doc = Document()

        section = doc.sections[0]
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.space_after = Pt(10)

        h1 = doc.styles["Heading 1"]
        h1.font.name = "Calibri"
        h1.font.size = Pt(16)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        h1.paragraph_format.space_before = Pt(18)
        h1.paragraph_format.space_after = Pt(12)

        h2 = doc.styles["Heading 2"]
        h2.font.name = "Calibri"
        h2.font.size = Pt(13)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
        h2.paragraph_format.space_before = Pt(14)
        h2.paragraph_format.space_after = Pt(8)

        self._render_header_and_letter(
            doc, title, client_name, client_position, client_entity, client_city, letter_text
        )

        section_num = 1
        if _has_content(context_text):
            doc.add_heading(f"{section_num}. CONTEXTO", level=1)
            _add_html_paragraphs(doc, context_text)
            section_num += 1

        for payload in schemes_payload:
            scheme_type = payload.get("scheme_type", "licensing")
            heading = self.SCHEME_HEADING.get(scheme_type, scheme_type.upper())
            self._render_scheme_block(doc, section_num, payload, products, scheme_heading=heading)
            section_num += 1

        self._render_legal_footer(doc, section_num, client_entity)

        footer = section.footer
        self._add_page_number(footer.paragraphs[0])
        return doc

    def generate_technical_annex(
        self, products: List[PortfolioProduct]
    ) -> Document:
        """
        Genera el anexo técnico con las descripciones de los productos.

        Returns:
            Objeto Document de python-docx.
        """
        doc = Document()

        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        doc.add_heading("ANEXO TÉCNICO", level=0)
        doc.add_paragraph(
            "Descripción detallada de las soluciones y componentes "
            "incluidos en el alcance de la propuesta comercial."
        )

        for i, product in enumerate(products, 1):
            doc.add_heading(f"{i}. {product.name}", level=1)
            doc.add_paragraph(f"Tipo: {product.product_type}")
            doc.add_paragraph("")
            if product.description:
                doc.add_paragraph(product.description)
            else:
                doc.add_paragraph("[Descripción pendiente de completar]")

        return doc

    def save_document(self, doc: Document, output_path: str) -> str:
        """Guarda el documento Word en la ruta especificada."""
        path = Path(output_path)
        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(path))
        return str(path)

    def convert_docx_to_pdf(self, docx_path: str, pdf_path: str) -> str:
        """
        Convierte un archivo .docx a PDF.
        En Windows usa docx2pdf (requiere Word).
        En otros sistemas intenta usar LibreOffice.

        Args:
            docx_path: Ruta al archivo .docx
            pdf_path: Ruta donde se guardará el PDF

        Returns:
            Ruta del PDF generado.
        """
        import os
        import platform

        docx_file = Path(docx_path)
        if not docx_file.exists():
            raise DocumentGeneratorError(f"El archivo no existe: {docx_path}")

        try:
            if platform.system() == "Windows":
                # docx2pdf usa win32com (COM de Windows) para automatizar Word.
                # FastAPI / Uvicorn ejecuta los endpoints síncronos en hilos del pool;
                # esos hilos no tienen COM inicializado, de ahí el error
                # "CoInitialize has not been called" (-2147221008).
                # Solución: inicializar y liberar COM explícitamente en este hilo.
                import pythoncom
                from docx2pdf import convert
                pythoncom.CoInitialize()
                try:
                    convert(docx_path, pdf_path)
                finally:
                    pythoncom.CoUninitialize()
            else:
                import subprocess
                pdf_dir = Path(pdf_path).parent
                pdf_dir.mkdir(parents=True, exist_ok=True)

                result = subprocess.run(
                    [
                        "libreoffice",
                        "--headless",
                        "--convert-to", "pdf",
                        "--outdir", str(pdf_dir),
                        str(docx_path)
                    ],
                    capture_output=True,
                    text=True,
                    timeout=60
                )

                if result.returncode != 0:
                    raise DocumentGeneratorError(
                        f"LibreOffice error: {result.stderr}"
                    )

                generated_pdf = docx_file.with_suffix(".pdf")

                import shutil
                if generated_pdf.exists() and str(generated_pdf) != str(pdf_path):
                    shutil.move(str(generated_pdf), str(pdf_path))

            if not Path(pdf_path).exists():
                raise DocumentGeneratorError(
                    "El PDF no fue generado correctamente"
                )

            return str(pdf_path)

        except Exception as e:
            raise DocumentGeneratorError(f"Error al convertir DOCX a PDF: {str(e)}")
