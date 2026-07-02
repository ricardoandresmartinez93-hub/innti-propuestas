"""
Pruebas unitarias para el generador de documentos Word.

Cubre:
  - _strip_html: limpieza de contenido HTML de TipTap
  - _has_content: detección de contenido real en HTML (Tarea 2)
  - generate_proposal_docx: lógica de fallback en secciones 5 y 6.1
  - _add_table_of_contents: campo TOC real con OOXML field codes
  - _build_proposal_docx (router): integración con campos de BD
"""
import io
import zipfile
import pytest
from pathlib import Path
from unittest.mock import MagicMock, patch

from app.services.document_generator import (
    DocumentGenerator,
    _add_html_paragraphs,
    _has_content,
    _strip_html,
)
from app.services.portfolio_service import PortfolioProduct


# ---------------------------------------------------------------------------
# Helpers de test
# ---------------------------------------------------------------------------

def _make_product(name: str = "Qx-Producto Test", category: str = "") -> PortfolioProduct:
    return PortfolioProduct(
        name=name,
        product_type="Plataforma",
        description="Descripción de prueba",
        business_framework="",
        revenue_info="",
        operational_costs="",
        monetization_model="",
        pricing_model="",
        country="Colombia",
        category=category,
    )


def _extract_text(doc) -> str:
    """Extrae todo el texto plano de un documento python-docx."""
    return "\n".join(p.text for p in doc.paragraphs)


def _build_minimal_docx(
    excluded_services: str | None = None,
    ip_section: str | None = None,
) -> object:
    """Invoca el generador con los parámetros mínimos necesarios."""
    gen = DocumentGenerator()
    return gen.generate_proposal_docx(
        title="Propuesta de Prueba",
        client_name="Juan Pérez",
        client_position="Gerente TI",
        client_entity="Entidad Prueba",
        client_city="Bogotá",
        scheme_types=["licensing"],
        products=[_make_product()],
        context_text="Contexto de prueba",
        scope_text="Alcance de prueba",
        letter_text="Carta de prueba",
        validity_period=None,
        economic_conditions=None,
        payment_terms=None,
        excluded_services=excluded_services,
        ip_section=ip_section,
    )


# ---------------------------------------------------------------------------
# Helper compartido: extrae el XML de word/document.xml del .docx generado
# ---------------------------------------------------------------------------

def _get_doc_xml(doc) -> str:
    """Serializa el Document a bytes y devuelve el XML interno de word/document.xml."""
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    with zipfile.ZipFile(buf) as z:
        return z.read("word/document.xml").decode("utf-8")


# ---------------------------------------------------------------------------
# Tests de _strip_html
# ---------------------------------------------------------------------------

class TestStripHtml:
    """Pruebas para la función utilitaria _strip_html."""

    def test_removes_paragraph_tags(self):
        result = _strip_html("<p>Texto de prueba</p>")
        assert "<p>" not in result
        assert "Texto de prueba" in result

    def test_removes_strong_and_em(self):
        result = _strip_html("<strong>Negrita</strong> y <em>cursiva</em>")
        assert "<strong>" not in result
        assert "Negrita" in result
        assert "cursiva" in result

    def test_decodes_html_entities(self):
        result = _strip_html("Tom&amp;Jerry &lt;personaje&gt; &quot;caricatura&quot;")
        assert "&amp;" not in result
        assert "Tom&Jerry" in result
        assert "<personaje>" in result
        assert '"caricatura"' in result

    def test_nbsp_becomes_space(self):
        result = _strip_html("Palabra&nbsp;separada")
        assert "&nbsp;" not in result
        assert "Palabra separada" in result

    def test_list_items_separated_by_newlines(self):
        html = "<ul><li>Ítem uno</li><li>Ítem dos</li></ul>"
        result = _strip_html(html)
        assert "Ítem uno" in result
        assert "Ítem dos" in result

    def test_br_becomes_newline(self):
        result = _strip_html("Línea A<br/>Línea B")
        assert "Línea A" in result
        assert "Línea B" in result

    def test_plain_text_unchanged(self):
        plain = "Texto sin etiquetas"
        assert _strip_html(plain) == plain

    def test_empty_string_returns_empty(self):
        assert _strip_html("") == ""

    def test_strips_nested_tags(self):
        html = "<div><p><strong>Título</strong></p><p>Párrafo</p></div>"
        result = _strip_html(html)
        assert "Título" in result
        assert "Párrafo" in result
        assert "<" not in result


# ---------------------------------------------------------------------------
# Tests de _has_content (Tarea 2 — helper de detección de contenido real)
# ---------------------------------------------------------------------------

class TestHasContent:
    """Tests para la función utilitaria _has_content."""

    def test_none_returns_false(self):
        assert not _has_content(None)

    def test_empty_string_returns_false(self):
        assert not _has_content("")

    def test_whitespace_only_returns_false(self):
        assert not _has_content("   \n  ")

    def test_empty_p_tag_returns_false(self):
        """TipTap genera <p></p> en pestañas vacías; no debe considerarse contenido."""
        assert not _has_content("<p></p>")

    def test_p_with_content_returns_true(self):
        assert _has_content("<p>Texto real</p>")

    def test_plain_text_returns_true(self):
        assert _has_content("Texto sin tags")

    def test_ul_with_items_returns_true(self):
        assert _has_content("<ul><li>Item</li></ul>")

    def test_nested_empty_tags_returns_false(self):
        """Tags anidados sin texto real no cuentan como contenido."""
        assert not _has_content("<p><br></p>")

    def test_strong_with_text_returns_true(self):
        assert _has_content("<strong>Negrita</strong>")


# ---------------------------------------------------------------------------
# Tests de _add_table_of_contents — campo TOC real con OOXML field codes
# ---------------------------------------------------------------------------

class TestTableOfContents:
    """Verifica que el documento contenga un campo TOC real, no texto placeholder."""

    def test_toc_field_code_is_present(self):
        """El XML del documento debe contener la instrucción del campo TOC."""
        doc = _build_minimal_docx()
        xml = _get_doc_xml(doc)
        assert "TOC" in xml, "No se encontró el campo TOC en el documento generado"

    def test_toc_has_dirty_true_for_auto_update(self):
        """El campo TOC debe llevar w:dirty='true' para que Word lo actualice al abrir."""
        doc = _build_minimal_docx()
        xml = _get_doc_xml(doc)
        # lxml puede serializar el atributo con comillas simples o dobles
        assert 'dirty="true"' in xml or "dirty='true'" in xml, (
            "El campo TOC no tiene el atributo w:dirty='true'"
        )

    def test_toc_includes_headings_1_to_3(self):
        """El campo TOC debe capturar los niveles Heading 1, 2 y 3."""
        doc = _build_minimal_docx()
        xml = _get_doc_xml(doc)
        assert "1-3" in xml, "El campo TOC no incluye la especificación de niveles 1-3"

    def test_toc_has_hyperlink_switch(self):
        """El campo TOC debe incluir el switch \\h para generar hipervínculos internos."""
        doc = _build_minimal_docx()
        xml = _get_doc_xml(doc)
        # La instrucción almacenada en instrText contiene los switches sin escape
        assert r"\h" in xml, "El campo TOC no incluye el switch \\h (hipervínculos)"

    def test_no_placeholder_text(self):
        """El texto estático de marcador de posición no debe aparecer en el documento."""
        doc = _build_minimal_docx()
        text = _extract_text(doc)
        assert "[El índice se genera automáticamente al actualizar campos en Word]" not in text

    def test_toc_title_is_present(self):
        """El encabezado 'TABLA DE CONTENIDO' debe seguir presente en el documento."""
        doc = _build_minimal_docx()
        text = _extract_text(doc)
        assert "TABLA DE CONTENIDO" in text


# ---------------------------------------------------------------------------
# Tests de generate_proposal_docx — sección 5 (Servicios Excluidos)
# ---------------------------------------------------------------------------

class TestSection5ExcludedServices:
    """Sección 5: el contenido de BD tiene precedencia sobre el texto hardcodeado."""

    def test_uses_db_content_when_provided(self):
        """Si excluded_services tiene contenido, debe aparecer en el documento."""
        doc = _build_minimal_docx(excluded_services="Servicio personalizado del cliente")
        text = _extract_text(doc)
        assert "Servicio personalizado del cliente" in text

    def test_db_content_replaces_hardcoded_list(self):
        """Cuando hay contenido en BD, el primer ítem hardcodeado NO debe aparecer."""
        hardcoded_first = DocumentGenerator.EXCLUDED_SERVICES[0]
        doc = _build_minimal_docx(excluded_services="Solo mi servicio excluido")
        text = _extract_text(doc)
        assert hardcoded_first not in text

    def test_fallback_to_hardcoded_when_none(self):
        """Sin contenido en BD (None), debe usar la lista hardcodeada completa."""
        hardcoded_first = DocumentGenerator.EXCLUDED_SERVICES[0]
        doc = _build_minimal_docx(excluded_services=None)
        text = _extract_text(doc)
        assert hardcoded_first in text

    def test_fallback_to_hardcoded_when_empty_string(self):
        """Cadena vacía: debe omitir la sección."""
        hardcoded_first = DocumentGenerator.EXCLUDED_SERVICES[0]
        doc = _build_minimal_docx(excluded_services="")
        text = _extract_text(doc)
        assert hardcoded_first not in text
        assert "SERVICIOS EXCLUIDOS" not in text

    def test_strips_html_from_db_content(self):
        """El contenido HTML de TipTap llega limpio al Word."""
        html_content = "<p>Servicio <strong>excluido</strong> del alcance</p>"
        doc = _build_minimal_docx(excluded_services=html_content)
        text = _extract_text(doc)
        assert "<p>" not in text
        assert "<strong>" not in text
        assert "Servicio excluido del alcance" in text


# ---------------------------------------------------------------------------
# Tests de generate_proposal_docx — sección 6.1 (Propiedad Intelectual)
# ---------------------------------------------------------------------------

class TestSection61IpSection:
    """Sección 6.1: el contenido de BD tiene precedencia sobre el texto hardcodeado."""

    def test_uses_db_content_when_provided(self):
        """Si ip_section tiene contenido, debe aparecer en el documento."""
        doc = _build_minimal_docx(ip_section="IP personalizada de la propuesta")
        text = _extract_text(doc)
        assert "IP personalizada de la propuesta" in text

    def test_db_content_replaces_hardcoded_ip_text(self):
        """Cuando hay contenido en BD, el texto hardcodeado de IP NO debe aparecer."""
        # El IP_TEXT hardcodeado empieza con "La arquitectura"
        doc = _build_minimal_docx(ip_section="Nuestros propios términos de IP")
        text = _extract_text(doc)
        assert "La arquitectura, los diseños técnicos" not in text

    def test_fallback_to_hardcoded_when_none(self):
        """Sin contenido en BD, debe usar el texto hardcodeado de IP."""
        doc = _build_minimal_docx(ip_section=None)
        text = _extract_text(doc)
        # El texto hardcodeado contiene "QUIPUX puede utilizarlos libremente"
        assert "QUIPUX puede utilizarlos libremente" in text

    def test_fallback_to_hardcoded_when_empty_string(self):
        """Cadena vacía: debe omitir la sección."""
        doc = _build_minimal_docx(ip_section="")
        text = _extract_text(doc)
        assert "Propiedad Intelectual" not in text

    def test_client_entity_interpolated_in_fallback(self):
        """El nombre del cliente se interpola correctamente en el fallback hardcodeado (cuando es None)."""
        gen = DocumentGenerator()
        doc = gen.generate_proposal_docx(
            title="Test",
            client_name="Ana López",
            client_position="Directora",
            client_entity="MiEntidad SA",
            client_city="Cali",
            scheme_types=["licensing"],
            products=[],
            context_text="",
            scope_text="",
            letter_text="",
            ip_section=None,
        )
        text = _extract_text(doc)
        assert "MiEntidad SA" in text

    def test_strips_html_from_db_content(self):
        """El contenido HTML de TipTap llega limpio al Word."""
        html_ip = "<p>Todos los derechos de <em>propiedad intelectual</em> pertenecen a QUIPUX.</p>"
        doc = _build_minimal_docx(ip_section=html_ip)
        text = _extract_text(doc)
        assert "<p>" not in text
        assert "<em>" not in text
        assert "propiedad intelectual" in text


# ---------------------------------------------------------------------------
# Test de integración: ambos campos activos simultáneamente
# ---------------------------------------------------------------------------

class TestBothFieldsTogether:
    """Verifica que ambos campos funcionen correctamente en simultáneo."""

    def test_both_db_fields_present(self):
        """Ambos campos de BD deben aparecer; ningún texto hardcodeado."""
        doc = _build_minimal_docx(
            excluded_services="Exclusión A",
            ip_section="IP Sección B",
        )
        text = _extract_text(doc)

        assert "Exclusión A" in text
        assert "IP Sección B" in text

        # Hardcoded NO debe aparecer
        assert DocumentGenerator.EXCLUDED_SERVICES[0] not in text
        assert "La arquitectura, los diseños técnicos" not in text

    def test_both_fields_empty_uses_all_hardcoded(self):
        """Con ambos campos None se usan los textos hardcodeados de fallback."""
        doc = _build_minimal_docx(excluded_services=None, ip_section=None)
        text = _extract_text(doc)

        assert DocumentGenerator.EXCLUDED_SERVICES[0] in text
        assert "Propiedad Intelectual" in text


# ---------------------------------------------------------------------------
# Test de integración con el router documents.py
# ---------------------------------------------------------------------------

class TestDocumentsRouterIntegration:
    """Verifica que _build_combined_docx propague el contenido por esquema al generador.

    El contenido por esquema vive en ProposalScheme; el router lo resuelve vía
    ``proposal_content_resolver`` y lo pasa al generador como ``schemes_payload``.
    """

    @staticmethod
    def _build_mock_scheme(
        scheme_type: str = "licensing",
        excluded_services=None,
        ip_section=None,
        scope_content="",
        economic_conditions=None,
        payment_terms=None,
        validity_period=None,
    ):
        mock_scheme = MagicMock()
        mock_scheme.id = 1
        mock_scheme.scheme_type.value = scheme_type
        mock_scheme.payment_frequency = "unico"
        mock_scheme.scope_content = scope_content
        mock_scheme.validity_period = validity_period
        mock_scheme.economic_conditions = economic_conditions
        mock_scheme.payment_terms = payment_terms
        mock_scheme.excluded_services = excluded_services
        mock_scheme.ip_section = ip_section
        return mock_scheme

    @staticmethod
    def _build_mock_proposal(scheme):
        mock_client = MagicMock()
        mock_client.name = "Ana García"
        mock_client.position = "Gerente"
        mock_client.entity = "Entidad Test"
        mock_client.city = "Medellín"

        mock_product = MagicMock()
        mock_product.product_name = "Producto Test"
        mock_product.category = ""

        mock_proposal = MagicMock()
        mock_proposal.id = 99
        mock_proposal.title = "Propuesta Test"
        mock_proposal.cover_title = None
        mock_proposal.client = mock_client
        mock_proposal.schemes = [scheme]
        mock_proposal.products = [mock_product]
        mock_proposal.context_content = ""
        mock_proposal.letter_content = ""
        mock_proposal.combine_schemes = True
        # Estos tests modelan el flujo legado (esquemas sin product_id)
        mock_proposal.uses_product_schemes = False
        return mock_proposal

    def test_router_passes_excluded_services_to_generator(self):
        """Si ProposalScheme.excluded_services tiene contenido, debe llegar al generador."""
        from app.routers.documents import _build_combined_docx
        import tempfile
        from pathlib import Path

        scheme = self._build_mock_scheme(
            excluded_services="Servicio excluido de prueba",
            ip_section="IP de prueba",
        )
        mock_proposal = self._build_mock_proposal(scheme)

        mock_db = MagicMock()
        mock_db.query.return_value.filter.return_value.first.return_value = mock_proposal

        mock_settings = MagicMock()
        mock_settings.portfolio_file_path = "dummy.xlsx"

        mock_portfolio = MagicMock()
        mock_portfolio.get_by_names.return_value = [_make_product()]

        with patch("app.routers.documents.PortfolioService", return_value=mock_portfolio):
            with patch("app.routers.documents.DocumentGenerator") as MockGenerator:
                mock_gen_instance = MagicMock()
                MockGenerator.return_value = mock_gen_instance
                mock_gen_instance.generate_combined_proposal_docx.return_value = MagicMock()
                mock_gen_instance.save_document.return_value = str(
                    Path(tempfile.gettempdir()) / "innti_docs" / "propuesta_99.docx"
                )

                _build_combined_docx(
                    proposal_id=99,
                    use_innti=False,
                    db=mock_db,
                    settings=mock_settings,
                )

                call_kwargs = mock_gen_instance.generate_combined_proposal_docx.call_args.kwargs
                schemes_payload = call_kwargs["schemes_payload"]
                assert len(schemes_payload) == 1
                assert schemes_payload[0]["excluded_services_text"] == "Servicio excluido de prueba"
                assert schemes_payload[0]["ip_section_text"] == "IP de prueba"

    def test_router_applies_default_ip_per_scheme_when_db_is_none(self):
        """Cuando ip_section / excluded_services son None en BD, el resolver aplica defaults por esquema.

        Para SaaS (services), excluded_services_text default es "" (no se renderiza la sección).
        Para licensing, debe aplicar el IP_TEXT_BY_SCHEME['licensing'] y la lista de exclusiones.
        """
        from app.routers.documents import _build_combined_docx
        from app.services.document_generator import DocumentGenerator
        import tempfile
        from pathlib import Path

        scheme = self._build_mock_scheme(
            scheme_type="services",
            excluded_services=None,
            ip_section=None,
        )
        mock_proposal = self._build_mock_proposal(scheme)

        mock_db = MagicMock()
        mock_db.query.return_value.filter.return_value.first.return_value = mock_proposal
        mock_settings = MagicMock()
        mock_settings.portfolio_file_path = "dummy.xlsx"

        mock_portfolio = MagicMock()
        mock_portfolio.get_by_names.return_value = []

        with patch("app.routers.documents.PortfolioService", return_value=mock_portfolio):
            with patch("app.routers.documents.DocumentGenerator") as MockGenerator:
                mock_gen_instance = MagicMock()
                MockGenerator.return_value = mock_gen_instance
                mock_gen_instance.generate_combined_proposal_docx.return_value = MagicMock()
                mock_gen_instance.save_document.return_value = str(
                    Path(tempfile.gettempdir()) / "innti_docs" / "propuesta_77.docx"
                )

                _build_combined_docx(
                    proposal_id=77,
                    use_innti=False,
                    db=mock_db,
                    settings=mock_settings,
                )

                call_kwargs = mock_gen_instance.generate_combined_proposal_docx.call_args.kwargs
                schemes_payload = call_kwargs["schemes_payload"]
                assert len(schemes_payload) == 1
                # SaaS: lista de exclusiones vacía → texto vacío
                assert schemes_payload[0]["excluded_services_text"] == ""
                # IP fallback: debe aplicar el texto de IP_TEXT_BY_SCHEME['services']
                ip_html = schemes_payload[0]["ip_section_text"]
                # El texto de SaaS empieza con "Los servicios prestados por QUIPUX a {client_entity}"
                assert "Los servicios prestados por QUIPUX" in ip_html
                assert "Entidad Test" in ip_html
                # Debe ser distinto del IP de licensing (que habla de "arquitectura")
                assert "arquitectura" not in ip_html.lower()


# ---------------------------------------------------------------------------
# Tests: convert_docx_to_pdf — inicialización COM en Windows
# ---------------------------------------------------------------------------

class TestConvertDocxToPdfWindowsCOM:
    """
    Verifica que convert_docx_to_pdf inicializa y libera COM correctamente en
    Windows para evitar el error 'CoInitialize has not been called'
    (-2147221008) que ocurre cuando FastAPI/Uvicorn ejecuta el endpoint en
    un hilo del pool que no tiene COM inicializado.
    """

    def test_com_initialize_called_on_windows(self, tmp_path):
        """CoInitialize y CoUninitialize deben llamarse en el path de Windows."""
        docx_file = tmp_path / "test.docx"
        docx_file.write_bytes(b"PK")  # contenido mínimo — docx2pdf se mockea
        pdf_path = str(tmp_path / "test.pdf")

        mock_pythoncom = MagicMock()

        # El mock crea el archivo PDF para pasar la validación de existencia
        def fake_convert(src, dst):
            Path(dst).write_bytes(b"%PDF-1.4")

        with patch("platform.system", return_value="Windows"), \
             patch.dict("sys.modules", {"pythoncom": mock_pythoncom}), \
             patch("docx2pdf.convert", side_effect=fake_convert):
            gen = DocumentGenerator()
            gen.convert_docx_to_pdf(str(docx_file), pdf_path)

        mock_pythoncom.CoInitialize.assert_called_once()
        mock_pythoncom.CoUninitialize.assert_called_once()

    def test_com_uninitialize_called_even_if_convert_raises(self, tmp_path):
        """CoUninitialize debe llamarse aunque docx2pdf.convert lance excepción."""
        docx_file = tmp_path / "test.docx"
        docx_file.write_bytes(b"PK")

        mock_pythoncom = MagicMock()

        with patch("platform.system", return_value="Windows"), \
             patch.dict("sys.modules", {"pythoncom": mock_pythoncom}), \
             patch("docx2pdf.convert", side_effect=RuntimeError("Word no disponible")):
            gen = DocumentGenerator()
            with pytest.raises(Exception):
                gen.convert_docx_to_pdf(str(docx_file), str(tmp_path / "out.pdf"))

        mock_pythoncom.CoInitialize.assert_called_once()
        mock_pythoncom.CoUninitialize.assert_called_once()


# ---------------------------------------------------------------------------
# Tests del parser HTML → docx (preservación de formato TipTap)
# ---------------------------------------------------------------------------


def _render(html: str) -> str:
    """Crea un Document, le pasa el HTML y devuelve el XML interno."""
    from docx import Document
    doc = Document()
    _add_html_paragraphs(doc, html)
    return _get_doc_xml(doc)


class TestHtmlToDocx:
    """Comprueba que cada marca/etiqueta HTML emitida por TipTap se traduce
    a la marca OOXML correspondiente en el .docx generado."""

    def test_bold_marker(self):
        xml = _render("<p><strong>negrita</strong></p>")
        assert "<w:b/>" in xml
        assert "negrita" in xml

    def test_italic_marker(self):
        xml = _render("<p><em>cursiva</em></p>")
        assert "<w:i/>" in xml
        assert "cursiva" in xml

    def test_underline_marker(self):
        xml = _render("<p><u>subrayado</u></p>")
        assert '<w:u w:val="single"/>' in xml

    def test_strike_marker(self):
        xml = _render("<p><s>tachado</s></p>")
        assert "<w:strike/>" in xml

    def test_superscript_marker(self):
        xml = _render("<p>x<sup>2</sup></p>")
        assert '<w:vertAlign w:val="superscript"/>' in xml

    def test_subscript_marker(self):
        xml = _render("<p>H<sub>2</sub>O</p>")
        assert '<w:vertAlign w:val="subscript"/>' in xml

    def test_text_align_center(self):
        xml = _render('<p style="text-align: center">centrado</p>')
        assert '<w:jc w:val="center"/>' in xml

    def test_text_align_right(self):
        xml = _render('<p style="text-align: right">derecha</p>')
        assert '<w:jc w:val="right"/>' in xml

    def test_text_align_justify(self):
        xml = _render('<p style="text-align: justify">justificado</p>')
        # python-docx serializa justify como "both"
        assert '<w:jc w:val="both"/>' in xml or '<w:jc w:val="justify"/>' in xml

    def test_blockquote_uses_quote_style(self):
        xml = _render("<blockquote>cita</blockquote>")
        assert 'w:val="Quote"' in xml

    def test_horizontal_rule(self):
        xml = _render("<hr>")
        assert "<w:pBdr" in xml
        assert "<w:bottom" in xml

    def test_hyperlink_creates_w_hyperlink(self):
        xml = _render('<p><a href="https://example.com">visita</a></p>')
        assert "<w:hyperlink" in xml
        assert "visita" in xml

    def test_highlight_marker(self):
        xml = _render('<p><mark style="background-color: #ffff00">resaltado</mark></p>')
        assert "<w:highlight" in xml

    def test_color_from_span(self):
        xml = _render('<p><span style="color: #ff0000">rojo</span></p>')
        assert '<w:color w:val="FF0000"' in xml

    def test_bullet_list(self):
        xml = _render("<ul><li>uno</li><li>dos</li></ul>")
        # python-docx usa "ListBullet" como ID interno del estilo "List Bullet"
        assert "ListBullet" in xml
        assert "uno" in xml
        assert "dos" in xml

    def test_ordered_list(self):
        xml = _render("<ol><li>a</li><li>b</li></ol>")
        assert "ListNumber" in xml

    def test_headings_use_heading_styles(self):
        xml = _render("<h1>Título 1</h1><h2>Título 2</h2><h3>Título 3</h3>")
        assert 'w:val="Heading1"' in xml
        assert 'w:val="Heading2"' in xml
        assert 'w:val="Heading3"' in xml

    def test_mixed_run_styles(self):
        """Un solo párrafo con runs de distintos formatos preserva cada uno."""
        xml = _render(
            "<p><strong>negrita</strong> normal <em>cursiva</em></p>"
        )
        assert "negrita" in xml
        assert "normal" in xml
        assert "cursiva" in xml
        assert "<w:b/>" in xml
        assert "<w:i/>" in xml

    def test_nested_marks(self):
        """Negrita + cursiva combinadas en un mismo texto."""
        xml = _render("<p><strong><em>ambas</em></strong></p>")
        assert "<w:b/>" in xml
        assert "<w:i/>" in xml
        assert "ambas" in xml

    def test_plain_text_without_tags(self):
        """Texto sin etiquetas se inserta como párrafo plano."""
        xml = _render("texto sin etiquetas")
        assert "texto sin etiquetas" in xml

    def test_empty_string_does_nothing(self):
        from docx import Document
        doc = Document()
        before = _get_doc_xml(doc)
        _add_html_paragraphs(doc, "")
        after = _get_doc_xml(doc)
        assert before == after

    def test_br_creates_line_break(self):
        xml = _render("<p>línea 1<br/>línea 2</p>")
        assert "<w:br/>" in xml
        assert "línea 1" in xml
        assert "línea 2" in xml

    def test_html_entities_decoded(self):
        xml = _render("<p>Tom&amp;Jerry &lt;eq&gt;</p>")
        # XML escapa & como &amp;, < como &lt;
        assert "Tom&amp;Jerry" in xml
        assert "&lt;eq&gt;" in xml

    def test_link_with_bold_inside(self):
        """Una marca dentro de un link debe coexistir con el hyperlink."""
        xml = _render('<p><a href="https://x.com"><strong>link bold</strong></a></p>')
        assert "<w:hyperlink" in xml
        assert "link bold" in xml
        assert "<w:b/>" in xml
