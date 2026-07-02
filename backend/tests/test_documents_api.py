"""
Tests para el router de generación de documentos (app/routers/documents.py).

Estrategia:
- Casos 404: propuesta inexistente → no requieren mocking.
- Casos de éxito: se parchea PortfolioService para evitar depender del xlsx.
  DocumentGenerator se ejecuta normalmente con python-docx.
- Error 500 en PDF: se parchean _build_proposal_docx y DocumentGenerator.
- Documentos separados: se verifican respuestas ZIP cuando combine_schemes=False.
"""
import io
import tempfile
import zipfile
import pytest
from pathlib import Path
from unittest.mock import patch, MagicMock
from fastapi import status
from docx import Document as DocxReader

from app.services.innti_service import InntiServiceError


def _docx_text_from_zip_bytes(zip_bytes: bytes, name_keyword: str) -> str:
    """Extrae todo el texto plano del .docx cuyo nombre contiene name_keyword."""
    z = zipfile.ZipFile(io.BytesIO(zip_bytes))
    name = next(n for n in z.namelist() if name_keyword in n and n.endswith(".docx"))
    with z.open(name) as f:
        doc = DocxReader(io.BytesIO(f.read()))
    return "\n".join(p.text for p in doc.paragraphs)

DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


# ── Helpers ───────────────────────────────────────────────────────────────────
def _create_proposal(client, sample_client_data, sample_proposal_data, creator_headers) -> int:
    c_res = client.post("/api/clients/", json=sample_client_data)
    assert c_res.status_code == status.HTTP_201_CREATED
    client_id = c_res.json()["id"]

    p_data = {**sample_proposal_data, "client_id": client_id}
    p_res = client.post("/api/proposals/", json=p_data, headers=creator_headers)
    assert p_res.status_code == status.HTTP_201_CREATED
    return p_res.json()["id"]


def _create_multi_product_proposal(client, sample_client_data, creator_headers) -> int:
    """Crea una propuesta con dos productos (cada uno con su esquema) y combine_schemes=False."""
    c_res = client.post("/api/clients/", json=sample_client_data)
    assert c_res.status_code == status.HTTP_201_CREATED
    client_id = c_res.json()["id"]

    p_data = {
        "title": "Propuesta multi-producto",
        "code": "TEST-001",
        "client_id": client_id,
        "combine_schemes": False,
        "products": [
            {
                "product_name": "Producto Licencias",
                "product_type": "Plataforma",
                "scheme": {"scheme_type": "licensing", "payment_frequency": "unico"},
            },
            {
                "product_name": "Producto Servicios",
                "product_type": "Plataforma",
                "scheme": {"scheme_type": "services", "payment_frequency": "mensual"},
            },
        ],
    }
    p_res = client.post("/api/proposals/", json=p_data, headers=creator_headers)
    assert p_res.status_code == status.HTTP_201_CREATED
    return p_res.json()["id"]


def _create_three_product_proposal(client, sample_client_data, creator_headers) -> dict:
    """Crea una propuesta con 3 productos, cada uno con su esquema y contenido distintivo.

    Devuelve el dict de la propuesta creada (con scheme ids para que el test
    pueda editar el contenido posteriormente).
    """
    c_res = client.post("/api/clients/", json=sample_client_data)
    client_id = c_res.json()["id"]
    p_data = {
        "title": "Propuesta tres productos Medellín",
        "code": "TEST-MED",
        "client_id": client_id,
        "combine_schemes": False,
        "products": [
            {
                "product_name": "Producto Licencias",
                "product_type": "Plataforma",
                "scheme": {
                    "scheme_type": "licensing",
                    "payment_frequency": "unico",
                    "scope_content": "<p>ALCANCE-LICENSING-UNICO</p>",
                    "economic_conditions": "<p>VALOR-LIC-100M</p>",
                    "payment_terms": "<p>PAGO-LIC-50-50</p>",
                    "ip_section": "<p>IP-LIC-MARKER</p>",
                },
            },
            {
                "product_name": "Producto Servicios",
                "product_type": "Plataforma",
                "scheme": {
                    "scheme_type": "services",
                    "payment_frequency": "mensual",
                    "scope_content": "<p>ALCANCE-SERVICES-MENSUAL</p>",
                    "economic_conditions": "<p>VALOR-SRV-MENSUAL</p>",
                    "payment_terms": "<p>PAGO-SRV-MENSUAL</p>",
                    "ip_section": "<p>IP-SRV-MARKER</p>",
                },
            },
            {
                "product_name": "Producto Soporte",
                "product_type": "Plataforma",
                "scheme": {
                    "scheme_type": "support_maintenance",
                    "payment_frequency": "anual",
                    "scope_content": "<p>ALCANCE-SUPPORT-ANUAL</p>",
                    "economic_conditions": "<p>VALOR-SUP-ANUAL</p>",
                    "payment_terms": "<p>PAGO-SUP-ANUAL</p>",
                    "ip_section": "<p>IP-SUP-MARKER</p>",
                },
            },
        ],
    }
    p_res = client.post("/api/proposals/", json=p_data, headers=creator_headers)
    assert p_res.status_code == status.HTTP_201_CREATED
    return p_res.json()


# ── Tests: generate-document ─────────────────────────────────────────────────
def test_generate_document_not_found(client):
    """Generar documento con propuesta inexistente → 404."""
    response = client.post("/api/proposals/99999/generate-document?use_innti=false")
    assert response.status_code == status.HTTP_404_NOT_FOUND


def test_generate_document_success(client, creator_headers, sample_client_data, sample_proposal_data):
    """
    Genera documento Word correctamente.
    Se parchea PortfolioService para devolver lista vacía de productos.
    El DocumentGenerator genera un .docx real con python-docx.
    """
    pid = _create_proposal(client, sample_client_data, sample_proposal_data, creator_headers)

    with patch("app.routers.documents.PortfolioService") as MockPortfolio:
        MockPortfolio.return_value.get_by_names.return_value = []

        response = client.post(
            f"/api/proposals/{pid}/generate-document",
            params={"use_innti": False},
        )

    assert response.status_code == status.HTTP_200_OK
    assert len(response.content) > 0  # el body contiene el archivo


# ── Tests: generate-pdf ───────────────────────────────────────────────────────
def test_generate_pdf_not_found(client):
    """Generar PDF con propuesta inexistente → 404."""
    response = client.post("/api/proposals/99999/generate-pdf?use_innti=false")
    assert response.status_code == status.HTTP_404_NOT_FOUND


def test_generate_pdf_conversion_error_returns_500(client, creator_headers, sample_client_data, sample_proposal_data):
    """
    Si convert_docx_to_pdf lanza una excepción, el endpoint devuelve 500
    con el mensaje 'No se pudo generar el PDF'.
    """
    pid = _create_proposal(client, sample_client_data, sample_proposal_data, creator_headers)

    # Crear un archivo .docx falso para que _build_proposal_docx lo devuelva
    tmpdir = Path(tempfile.gettempdir()) / "innti_docs"
    tmpdir.mkdir(exist_ok=True)
    fake_docx = tmpdir / f"propuesta_{pid}.docx"
    fake_docx.write_bytes(b"PK\x03\x04")  # cabecera mínima de ZIP/DOCX

    with patch("app.routers.documents._build_combined_docx", return_value=fake_docx):
        with patch("app.routers.documents.DocumentGenerator") as MockDocGen:
            MockDocGen.return_value.convert_docx_to_pdf.side_effect = Exception(
                "WeasyPrint no disponible"
            )
            response = client.post(
                f"/api/proposals/{pid}/generate-pdf",
                params={"use_innti": False},
            )

    assert response.status_code == status.HTTP_500_INTERNAL_SERVER_ERROR
    assert "No se pudo generar el PDF" in response.json()["detail"]


# ── Tests: documentos separados (combine_schemes=False) ──────────────────────
def test_generate_document_separate_returns_zip_per_product(client, creator_headers, sample_client_data):
    """
    Con combine_schemes=False y 2 productos, generate-document devuelve un ZIP
    que contiene un .docx POR PRODUCTO (nombre con slug del producto).
    """
    pid = _create_multi_product_proposal(client, sample_client_data, creator_headers)

    with patch("app.routers.documents.PortfolioService") as MockPortfolio:
        MockPortfolio.return_value.get_by_names.return_value = []
        response = client.post(
            f"/api/proposals/{pid}/generate-document",
            params={"use_innti": False},
        )

    assert response.status_code == status.HTTP_200_OK
    assert response.headers["content-type"] == "application/zip"

    z = zipfile.ZipFile(io.BytesIO(response.content))
    names = z.namelist()
    assert len(names) == 2
    assert any("producto-licencias" in n for n in names)
    assert any("producto-servicios" in n for n in names)


def test_generate_pdf_separate_returns_zip(client, creator_headers, sample_client_data):
    """
    Con combine_schemes=False y 2 productos, generate-pdf devuelve un ZIP
    con un .pdf por producto (convert_docx_to_pdf se parchea para crear el .pdf vacío).
    """
    pid = _create_multi_product_proposal(client, sample_client_data, creator_headers)

    def fake_convert(docx_path: str, pdf_path: str) -> None:
        Path(pdf_path).write_bytes(b"%PDF-1.4")

    with patch("app.routers.documents.PortfolioService") as MockPortfolio:
        MockPortfolio.return_value.get_by_names.return_value = []
        with patch.object(
            __import__("app.services.document_generator", fromlist=["DocumentGenerator"]).DocumentGenerator,
            "convert_docx_to_pdf",
            side_effect=fake_convert,
        ):
            response = client.post(
                f"/api/proposals/{pid}/generate-pdf",
                params={"use_innti": False},
            )

    assert response.status_code == status.HTTP_200_OK
    assert response.headers["content-type"] == "application/zip"

    z = zipfile.ZipFile(io.BytesIO(response.content))
    names = z.namelist()
    assert len(names) == 2
    assert all(n.endswith(".pdf") for n in names)


def test_create_proposal_separate_with_single_product_rejected(
    client, creator_headers, sample_client_data
):
    """combine_schemes=False con un único producto debe ser rechazado con 422.

    No tiene sentido pedir "Documentos separados" si solo hay un producto —
    el validator de ProposalCreate lo rechaza explícitamente.
    """
    c_res = client.post("/api/clients/", json=sample_client_data)
    client_id = c_res.json()["id"]
    p_data = {
        "title": "Separado con un producto",
        "code": "TEST-SEP1",
        "client_id": client_id,
        "combine_schemes": False,
        "products": [{
            "product_name": "Producto Único",
            "product_type": "Plataforma",
            "scheme": {"scheme_type": "licensing", "payment_frequency": "unico"},
        }],
    }
    p_res = client.post("/api/proposals/", json=p_data, headers=creator_headers)
    assert p_res.status_code == status.HTTP_422_UNPROCESSABLE_ENTITY
    assert "combine_schemes" in p_res.text


def test_generate_document_combined_returns_docx(
    client, creator_headers, sample_client_data, sample_proposal_data
):
    """Con combine_schemes=True, generate-document devuelve un .docx único (nunca ZIP)."""
    pid = _create_proposal(client, sample_client_data, sample_proposal_data, creator_headers)

    with patch("app.routers.documents.PortfolioService") as MockPortfolio:
        MockPortfolio.return_value.get_by_names.return_value = []
        response = client.post(
            f"/api/proposals/{pid}/generate-document",
            params={"use_innti": False},
        )

    assert response.status_code == status.HTTP_200_OK
    assert "wordprocessingml" in response.headers["content-type"]


def test_generate_document_separate_innti_per_scheme(client, creator_headers, sample_client_data):
    """
    Con combine_schemes=False y use_innti=True, las secciones de forma de pago y
    condiciones económicas se generan con el tipo de esquema individual (no combinado).
    Cada documento debe recibir el contenido de pago correspondiente a su esquema.
    """
    pid = _create_multi_product_proposal(client, sample_client_data, creator_headers)

    with patch("app.routers.documents.PortfolioService") as MockPortfolio:
        MockPortfolio.return_value.get_by_names.return_value = []

        with patch("app.routers.documents.InntiService") as MockInntiCls:
            mock_innti = MagicMock()
            MockInntiCls.return_value = mock_innti
            mock_innti.generate_context_section.return_value = "context"
            mock_innti.generate_scope_section.return_value = "scope"
            mock_innti.generate_cover_letter.return_value = "letter"
            mock_innti.generate_validity_section.return_value = "validity"
            mock_innti.generate_economic_conditions_section.return_value = "economic"
            mock_innti.generate_payment_terms_section.return_value = "payment"
            mock_innti.generate_excluded_services_section.return_value = "excluded"
            mock_innti.generate_ip_section.return_value = "ip"

            response = client.post(
                f"/api/proposals/{pid}/generate-document",
                params={"use_innti": True},
            )

    assert response.status_code == status.HTTP_200_OK
    assert response.headers["content-type"] == "application/zip"

    # Verificar que generate_payment_terms_section fue llamada con cada esquema individual
    payment_calls = [
        call.args[0]
        for call in mock_innti.generate_payment_terms_section.call_args_list
    ]
    assert "licensing" in payment_calls, (
        f"Se esperaba llamada con 'licensing', se recibieron: {payment_calls}"
    )
    assert "services" in payment_calls, (
        f"Se esperaba llamada con 'services', se recibieron: {payment_calls}"
    )

    # Verificar que generate_economic_conditions_section fue llamada con cada esquema individual
    economic_calls = [
        call.args[1]  # segundo arg es scheme_type
        for call in mock_innti.generate_economic_conditions_section.call_args_list
    ]
    assert "licensing" in economic_calls
    assert "services" in economic_calls

    # El alcance de cada esquema se genera SOLO con su producto vinculado
    scope_calls = [call.args[0] for call in mock_innti.generate_scope_section.call_args_list]
    assert ["Producto Licencias"] in scope_calls
    assert ["Producto Servicios"] in scope_calls


# ── Tests: contenido diferenciado por producto (regresión del bug reportado) ──
def test_separate_zip_contains_different_content_per_product(
    client, creator_headers, sample_client_data
):
    """REGRESIÓN: el bug reportado era que el ZIP separado tenía N archivos
    con contenido idéntico (solo cambiaba el nombre). Ahora cada .docx debe
    incluir el contenido distintivo del esquema de SU producto (alcance, IP, pago).
    """
    proposal = _create_three_product_proposal(client, sample_client_data, creator_headers)
    pid = proposal["id"]

    with patch("app.routers.documents.PortfolioService") as MockPortfolio:
        MockPortfolio.return_value.get_by_names.return_value = []
        response = client.post(
            f"/api/proposals/{pid}/generate-document",
            params={"use_innti": False},
        )

    assert response.status_code == status.HTTP_200_OK
    assert response.headers["content-type"] == "application/zip"

    text_lic = _docx_text_from_zip_bytes(response.content, "producto-licencias")
    text_srv = _docx_text_from_zip_bytes(response.content, "producto-servicios")
    text_sup = _docx_text_from_zip_bytes(response.content, "producto-soporte")

    # Cada documento contiene SU propio alcance, no el de los otros
    assert "ALCANCE-LICENSING-UNICO" in text_lic
    assert "ALCANCE-SERVICES-MENSUAL" not in text_lic
    assert "ALCANCE-SUPPORT-ANUAL" not in text_lic

    assert "ALCANCE-SERVICES-MENSUAL" in text_srv
    assert "ALCANCE-LICENSING-UNICO" not in text_srv

    assert "ALCANCE-SUPPORT-ANUAL" in text_sup
    assert "ALCANCE-LICENSING-UNICO" not in text_sup

    # Cada documento contiene SU propio IP, no el de los otros
    assert "IP-LIC-MARKER" in text_lic
    assert "IP-SRV-MARKER" not in text_lic
    assert "IP-SRV-MARKER" in text_srv
    assert "IP-SUP-MARKER" in text_sup

    # Cada documento contiene SU propia forma de pago
    assert "PAGO-LIC-50-50" in text_lic
    assert "PAGO-SRV-MENSUAL" in text_srv
    assert "PAGO-SUP-ANUAL" in text_sup


def test_separate_zip_saas_omits_excluded_services_section(
    client, creator_headers, sample_client_data
):
    """SaaS (services) no debe renderizar la sección "SERVICIOS EXCLUIDOS"
    cuando no se ha provisto contenido — regla del PDF de la reunión.
    """
    c_res = client.post("/api/clients/", json=sample_client_data)
    client_id = c_res.json()["id"]
    p_data = {
        "title": "Propuesta SaaS + Licensing",
        "code": "TEST-SAAS",
        "client_id": client_id,
        "combine_schemes": False,
        "products": [
            {
                "product_name": "Producto Licencias",
                "product_type": "Plataforma",
                "scheme": {"scheme_type": "licensing", "payment_frequency": "unico"},
            },
            {
                "product_name": "Producto Servicios",
                "product_type": "Plataforma",
                "scheme": {"scheme_type": "services", "payment_frequency": "mensual"},
            },
        ],
    }
    p_res = client.post("/api/proposals/", json=p_data, headers=creator_headers)
    pid = p_res.json()["id"]

    with patch("app.routers.documents.PortfolioService") as MockPortfolio:
        MockPortfolio.return_value.get_by_names.return_value = []
        response = client.post(
            f"/api/proposals/{pid}/generate-document",
            params={"use_innti": False},
        )

    assert response.status_code == status.HTTP_200_OK
    text_lic = _docx_text_from_zip_bytes(response.content, "producto-licencias")
    text_srv = _docx_text_from_zip_bytes(response.content, "producto-servicios")

    # Licensing SÍ debe tener la sección de exclusiones (heading)
    assert "SERVICIOS EXCLUIDOS" in text_lic.upper()
    # SaaS NO debe tener esa sección
    assert "SERVICIOS EXCLUIDOS" not in text_srv.upper()


def test_combined_docx_contains_block_per_product(client, creator_headers, sample_client_data):
    """Modo unificado: un único .docx con un bloque «PRODUCTO — ESQUEMA» por producto."""
    proposal = _create_three_product_proposal(client, sample_client_data, creator_headers)
    pid = proposal["id"]
    # Cambiar a unificado vía PATCH
    client.patch(f"/api/proposals/{pid}", json={"combine_schemes": True})

    with patch("app.routers.documents.PortfolioService") as MockPortfolio:
        MockPortfolio.return_value.get_by_names.return_value = []
        response = client.post(
            f"/api/proposals/{pid}/generate-document",
            params={"use_innti": False},
        )

    assert response.status_code == status.HTTP_200_OK
    assert "wordprocessingml" in response.headers["content-type"]

    doc = DocxReader(io.BytesIO(response.content))
    text = "\n".join(p.text for p in doc.paragraphs)
    # Aparece un bloque por producto, titulado «PRODUCTO — ESQUEMA»
    assert "PRODUCTO LICENCIAS — LICENCIAMIENTO" in text
    assert "PRODUCTO SERVICIOS — PRESTACIÓN DE SERVICIOS" in text
    assert "PRODUCTO SOPORTE — SOPORTE Y MANTENIMIENTO" in text
    # Y los IP de cada esquema están todos presentes
    assert "IP-LIC-MARKER" in text
    assert "IP-SRV-MARKER" in text
    assert "IP-SUP-MARKER" in text


def test_combined_docx_two_products_same_scheme_two_blocks(
    client, creator_headers, sample_client_data
):
    """Dos productos con el MISMO tipo de esquema generan DOS bloques en el unificado."""
    c_res = client.post("/api/clients/", json=sample_client_data)
    client_id = c_res.json()["id"]
    p_data = {
        "title": "Dos licencias",
        "code": "TEST-2LIC",
        "client_id": client_id,
        "combine_schemes": True,
        "products": [
            {
                "product_name": "Producto Alfa",
                "product_type": "Plataforma",
                "scheme": {
                    "scheme_type": "licensing",
                    "payment_frequency": "unico",
                    "economic_conditions": "<p>VALOR-ALFA</p>",
                },
            },
            {
                "product_name": "Producto Beta",
                "product_type": "Plataforma",
                "scheme": {
                    "scheme_type": "licensing",
                    "payment_frequency": "unico",
                    "economic_conditions": "<p>VALOR-BETA</p>",
                },
            },
        ],
    }
    p_res = client.post("/api/proposals/", json=p_data, headers=creator_headers)
    pid = p_res.json()["id"]

    with patch("app.routers.documents.PortfolioService") as MockPortfolio:
        MockPortfolio.return_value.get_by_names.return_value = []
        response = client.post(
            f"/api/proposals/{pid}/generate-document",
            params={"use_innti": False},
        )

    assert response.status_code == status.HTTP_200_OK
    doc = DocxReader(io.BytesIO(response.content))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "PRODUCTO ALFA — LICENCIAMIENTO" in text
    assert "PRODUCTO BETA — LICENCIAMIENTO" in text
    assert "VALOR-ALFA" in text
    assert "VALOR-BETA" in text


# ── Tests: propuestas legadas (esquemas sin product_id) ──────────────────────
def _create_legacy_proposal(db_session, combine_schemes: bool):
    """Crea directamente en BD una propuesta legada: esquemas SIN product_id."""
    from app.models.client import Client
    from app.models.proposal import Proposal, ProposalScheme, SchemeType

    client_obj = Client(name="Cliente Legado", entity="Entidad Legada")
    db_session.add(client_obj)
    db_session.flush()

    proposal = Proposal(
        title="Propuesta legada",
        code="LEG-001",
        client_id=client_obj.id,
        combine_schemes=combine_schemes,
    )
    db_session.add(proposal)
    db_session.flush()

    db_session.add(ProposalScheme(
        proposal_id=proposal.id,
        scheme_type=SchemeType.LICENSING,
        payment_frequency="unico",
        ip_section="<p>IP-LEGADA-LIC</p>",
    ))
    db_session.add(ProposalScheme(
        proposal_id=proposal.id,
        scheme_type=SchemeType.SERVICES,
        payment_frequency="mensual",
        ip_section="<p>IP-LEGADA-SRV</p>",
    ))
    db_session.commit()
    return proposal.id


def test_legacy_proposal_separate_keeps_per_scheme_files(
    client, db_session, creator_headers
):
    """Propuesta legada + separado: un archivo POR ESQUEMA (comportamiento anterior)."""
    pid = _create_legacy_proposal(db_session, combine_schemes=False)

    with patch("app.routers.documents.PortfolioService") as MockPortfolio:
        MockPortfolio.return_value.get_by_names.return_value = []
        response = client.post(
            f"/api/proposals/{pid}/generate-document",
            params={"use_innti": False},
        )

    assert response.status_code == status.HTTP_200_OK
    assert response.headers["content-type"] == "application/zip"
    z = zipfile.ZipFile(io.BytesIO(response.content))
    names = z.namelist()
    assert len(names) == 2
    assert any("licenciamiento" in n for n in names)
    assert any("servicios" in n for n in names)


def test_legacy_proposal_combined_keeps_scheme_headings(
    client, db_session, creator_headers
):
    """Propuesta legada + unificado: bloques «ESQUEMA: …» (comportamiento anterior)."""
    pid = _create_legacy_proposal(db_session, combine_schemes=True)

    with patch("app.routers.documents.PortfolioService") as MockPortfolio:
        MockPortfolio.return_value.get_by_names.return_value = []
        response = client.post(
            f"/api/proposals/{pid}/generate-document",
            params={"use_innti": False},
        )

    assert response.status_code == status.HTTP_200_OK
    doc = DocxReader(io.BytesIO(response.content))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "ESQUEMA: LICENCIAMIENTO" in text
    assert "ESQUEMA: PRESTACIÓN DE SERVICIOS" in text
    assert "IP-LEGADA-LIC" in text
    assert "IP-LEGADA-SRV" in text


# ── Tests: generate-annex ────────────────────────────────────────────────────
def test_generate_annex_not_found(client):
    """Generar anexo técnico con propuesta inexistente → 404."""
    response = client.post("/api/proposals/99999/generate-annex")
    assert response.status_code == status.HTTP_404_NOT_FOUND


def test_generate_annex_success(client, creator_headers, sample_client_data, sample_proposal_data):
    """
    Genera el anexo técnico correctamente.
    Se parchea PortfolioService; DocumentGenerator se ejecuta con python-docx.
    """
    pid = _create_proposal(client, sample_client_data, sample_proposal_data, creator_headers)

    with patch("app.routers.documents.PortfolioService") as MockPortfolio:
        MockPortfolio.return_value.get_by_names.return_value = []

        response = client.post(f"/api/proposals/{pid}/generate-annex")

    assert response.status_code == status.HTTP_200_OK
    assert len(response.content) > 0


# ── Tests: resiliencia ante fallos parciales de Innti ─────────────────────────
def test_generate_document_innti_cover_letter_failure_uses_fallback(
    client, creator_headers, sample_client_data, sample_proposal_data
):
    """
    Cuando generate_cover_letter lanza InntiServiceError se usa el template
    de respaldo: letter_content queda NO vacío y las demás secciones se persisten
    (las globales en Proposal, las por esquema en ProposalScheme).
    """
    pid = _create_proposal(client, sample_client_data, sample_proposal_data, creator_headers)

    with patch("app.routers.documents.PortfolioService") as MockPortfolio:
        MockPortfolio.return_value.get_by_names.return_value = []

        with patch("app.routers.documents.InntiService") as MockInntiCls:
            mock_innti = MagicMock()
            MockInntiCls.return_value = mock_innti
            mock_innti.generate_context_section.return_value = "contexto generado"
            mock_innti.generate_scope_section.return_value = "alcance generado"
            mock_innti.generate_cover_letter.side_effect = InntiServiceError("timeout")
            mock_innti.generate_validity_section.return_value = "plazo generado"
            mock_innti.generate_economic_conditions_section.return_value = "condiciones"
            mock_innti.generate_payment_terms_section.return_value = "forma de pago"

            response = client.post(
                f"/api/proposals/{pid}/generate-document",
                params={"use_innti": True},
            )

    assert response.status_code == status.HTTP_200_OK

    updated = client.get(f"/api/proposals/{pid}")
    data = updated.json()
    assert data["context_content"] == "contexto generado", (
        "context_content (global) debe persistirse aunque cover_letter falle"
    )
    # Con el fallback, letter_content debe ser el template (no vacío)
    assert data["letter_content"], (
        "letter_content debe usar el template de respaldo cuando Innti falla"
    )
    assert "Juan Pablo Ramírez Madrid" in data["letter_content"], (
        "El template de respaldo debe incluir la firma del VP"
    )
    # El contenido por esquema vive ahora en ProposalScheme
    assert data["schemes"][0]["scope_content"] == "alcance generado", (
        "scope_content debe persistirse en el esquema"
    )
    assert data["schemes"][0]["payment_terms"] == "forma de pago"


def test_generate_document_innti_cover_letter_empty_uses_fallback(
    client, creator_headers, sample_client_data, sample_proposal_data
):
    """
    Cuando generate_cover_letter devuelve string vacío (sin excepción)
    se usa igualmente el template de respaldo.
    """
    pid = _create_proposal(client, sample_client_data, sample_proposal_data, creator_headers)

    with patch("app.routers.documents.PortfolioService") as MockPortfolio:
        MockPortfolio.return_value.get_by_names.return_value = []

        with patch("app.routers.documents.InntiService") as MockInntiCls:
            mock_innti = MagicMock()
            MockInntiCls.return_value = mock_innti
            mock_innti.generate_context_section.return_value = "contexto"
            mock_innti.generate_scope_section.return_value = "alcance"
            mock_innti.generate_cover_letter.return_value = ""   # vacío, no excepción
            mock_innti.generate_validity_section.return_value = "plazo"
            mock_innti.generate_economic_conditions_section.return_value = "condiciones"
            mock_innti.generate_payment_terms_section.return_value = "pago"
            mock_innti.generate_excluded_services_section.return_value = "excluidos"
            mock_innti.generate_ip_section.return_value = "ip"

            response = client.post(
                f"/api/proposals/{pid}/generate-document",
                params={"use_innti": True},
            )

    assert response.status_code == status.HTTP_200_OK

    updated = client.get(f"/api/proposals/{pid}")
    data = updated.json()
    assert data["letter_content"], (
        "letter_content debe usar el template cuando Innti devuelve vacío"
    )
    assert "Juan Pablo Ramírez Madrid" in data["letter_content"]


def test_generate_document_innti_all_sections_persisted(
    client, creator_headers, sample_client_data, sample_proposal_data
):
    """Con todas las llamadas Innti exitosas, secciones globales se persisten en Proposal
    y secciones por esquema se persisten en cada ProposalScheme.
    """
    pid = _create_proposal(client, sample_client_data, sample_proposal_data, creator_headers)

    with patch("app.routers.documents.PortfolioService") as MockPortfolio:
        MockPortfolio.return_value.get_by_names.return_value = []

        with patch("app.routers.documents.InntiService") as MockInntiCls:
            mock_innti = MagicMock()
            MockInntiCls.return_value = mock_innti
            mock_innti.generate_context_section.return_value = "contexto"
            mock_innti.generate_scope_section.return_value = "alcance"
            mock_innti.generate_cover_letter.return_value = "<p>Carta generada</p>"
            mock_innti.generate_validity_section.return_value = "plazo"
            mock_innti.generate_economic_conditions_section.return_value = "condiciones"
            mock_innti.generate_payment_terms_section.return_value = "pago"

            response = client.post(
                f"/api/proposals/{pid}/generate-document",
                params={"use_innti": True},
            )

    assert response.status_code == status.HTTP_200_OK

    updated = client.get(f"/api/proposals/{pid}")
    data = updated.json()
    assert data["context_content"] == "contexto"
    assert data["letter_content"] == "<p>Carta generada</p>"
    # El contenido por esquema vive ahora en ProposalScheme
    scheme = data["schemes"][0]
    assert scheme["scope_content"] == "alcance"
    assert scheme["validity_period"] == "plazo"
    assert scheme["economic_conditions"] == "condiciones"
    assert scheme["payment_terms"] == "pago"
